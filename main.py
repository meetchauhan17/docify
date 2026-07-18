import os, uuid, time, json, re, sys, base64, io
from pathlib import Path
from PIL import Image, ImageEnhance, ImageFilter, ImageOps
from fastapi import FastAPI, UploadFile, File, Form, Request
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
import pypdfium2 as pdfium
from docx import Document
from docx.shared import Pt, Cm, Inches, Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.pdfgen import canvas
from reportlab.lib.utils import ImageReader
from reportlab.lib.pagesizes import A4
from dotenv import load_dotenv
from google import genai
from google.genai import types as genai_types

load_dotenv(override=True)

# Force UTF-8 output so Greek/math symbols in OCR results don't crash the console
try:
    sys.stdout.reconfigure(encoding="utf-8", errors="replace")
except Exception:
    pass  # Python < 3.7 or non-TextIOWrapper stdout

# ── Multi-key Gemini rotation ──────────────────────
# Support comma-separated list: GEMINI_API_KEY=key1,key2,key3
_GEMINI_KEYS_RAW = os.getenv("GEMINI_API_KEY", "")
_GEMINI_KEYS: list[str] = [k.strip() for k in _GEMINI_KEYS_RAW.split(",") if k.strip()]
_gemini_key_idx = 0   # current key index (rotated on 429/quota)

def _get_gemini_key() -> str:
    return _GEMINI_KEYS[_gemini_key_idx % len(_GEMINI_KEYS)] if _GEMINI_KEYS else ""

def _rotate_gemini_key():
    global _gemini_key_idx
    _gemini_key_idx += 1
    new_key = _get_gemini_key()
    print(f"  [KEY-ROTATE] Switched to Gemini key index {_gemini_key_idx % max(1, len(_GEMINI_KEYS))}")
    return new_key

# Build initial Gemini client
def _make_gemini_client(api_key: str):
    return genai.Client(
        api_key=api_key,
        http_options=genai_types.HttpOptions(timeout=120_000),
    ) if api_key else None


def _classify_text_element(txt: str) -> tuple[bool, bool]:
    """
    General classifier for text elements in any document.
    Returns (is_note, is_metadata).
    """
    txt_strip = txt.strip()
    if not txt_strip:
        return False, False
    lower_txt = txt_strip.lower()

    # 1. Footnote/Signature/Note keywords
    note_start_keywords = {
        "note", "batch", "bacth", "prepared", "approved", "signature", "signed", "sign",
        "prof", "head", "director", "manager", "chairman", "dean", "officer", "coordinator",
        "clerk", "assistant", "footnote", "remark", "remarks", "instruction", "instructions"
    }
    note_contains_keywords = {
        "prof.", "dr.", "prepared by", "approved by", "asst. prof", "head, cse", "signature of",
        "authorized signatory", "seal & signature", "head of department"
    }

    first_word = lower_txt.split()[0] if lower_txt.split() else ""
    # Remove trailing punctuation
    first_word = "".join(char for char in first_word if char.isalnum())

    is_note = False
    if first_word in note_start_keywords:
        is_note = True
    elif any(kw in lower_txt for kw in note_contains_keywords):
        is_note = True

    # 2. Metadata keywords or Key-Value format
    is_metadata = False
    metadata_start_keywords = {
        "date", "page", "duration", "class", "year", "sem", "subject", "time", "location",
        "remarks", "author", "version", "title", "dept", "department", "id", "no", "doc", "document"
    }
    
    if not is_note:
        if first_word in metadata_start_keywords:
            is_metadata = True
        elif ":" in lower_txt:
            key_part = lower_txt.split(":")[0].strip()
            if (len(key_part) > 0 and len(key_part) <= 25 and 
                    all(char.isalnum() or char.isspace() or char in "-_/" for char in key_part)):
                is_metadata = True

    return is_note, is_metadata

# Groq & OpenRouter keys
GROQ_API_KEY       = os.getenv("GROQ_API_KEY", "")
OPENROUTER_API_KEY = os.getenv("OPENROUTER_API_KEY", "")

# Track which provider last succeeded (for status endpoint)
_last_provider: str = "none"
_provider_status: dict = {
    "gemini":     "unknown",
    "groq":       "unknown" if GROQ_API_KEY else "not_configured",
    "openrouter": "unknown" if OPENROUTER_API_KEY else "not_configured",
}

from pydantic import BaseModel, Field
from typing import List, Optional

class DocElement(BaseModel):
    type: str = Field(
        ...,
        description="The type of the element. Must be one of: 'text', 'blank_line', 'table', 'drawing'."
    )
    text: Optional[str] = Field(
        None,
        description="The transcribed text content of the line or paragraph. For multi-line content within a single element, use \\n to separate lines."
    )
    tag: Optional[str] = Field(
        None,
        description="The formatting tag for text elements: 'HEADING' (major title/heading), 'SUBHEAD' (sub-heading/section label), 'BODY' (normal body paragraph), 'BULLET' (bullet point or numbered list item), 'CENTER' (visually centered line), 'UNDERLN' (underlined text)."
    )
    bold: Optional[bool] = Field(
        None,
        description="True if the text is bold or written significantly darker/thicker than normal."
    )
    italic: Optional[bool] = Field(
        None,
        description="True if the text is italicized or slanted."
    )
    underline: Optional[bool] = Field(
        None,
        description="True if the text is underlined."
    )
    alignment: Optional[str] = Field(
        None,
        description="Text alignment: 'left', 'center', 'right', 'justify'."
    )
    font_size_pt: Optional[float] = Field(
        None,
        description="Estimated font size in points (pt). Use 16 for main headings, 13 for sub-headings, 11-12 for body text, 10 for captions/footnotes. Only set when the size visually differs from normal body text."
    )
    left_indent_cm: Optional[float] = Field(
        None,
        description="Estimated left indentation of the text element in centimeters (usually 0.0, or 0.5 for bullet points/indented blocks)."
    )
    table_data: Optional[List[List[str]]] = Field(
        None,
        description="A list of rows, where each row is a list of strings representing the cells of the table. Each cell string must contain the FULL text of that cell. Use \\n within a cell string for multi-line cell content."
    )
    borderless: Optional[bool] = Field(
        None,
        description="True if the table is a layout grid (e.g. side-by-side signature blocks, multi-column key-value pairs) and should be rendered without visible borders."
    )
    col_alignments: Optional[List[str]] = Field(
        None,
        description="Per-column text alignment for table columns. Each entry should be 'left', 'center', or 'right'. Length must match the number of columns in table_data."
    )
    bbox: Optional[List[float]] = Field(
        None,
        description="Normalized bounding box [x1, y1, x2, y2] (from 0.0 to 1.0) of the drawing or diagram in the image. "
                    "The box MUST be drawn TIGHTLY around the visible ink/marks only — do NOT pad it with surrounding whitespace. "
                    "x1,y1 is top-left corner; x2,y2 is bottom-right corner of the drawn marks."
    )
    description: Optional[str] = Field(
        None,
        description="A one-sentence description of the drawing, diagram, sketch, or graph."
    )
    is_simple_arrow: Optional[bool] = Field(
        None,
        description="True if the entire drawing is just a single plain hand-drawn arrow (no labels, text, or boxes)."
    )
    arrow_direction: Optional[str] = Field(
        None,
        description="If is_simple_arrow is True, the direction the arrow points: 'right', 'left', 'up', 'down', 'up-right', 'up-left', 'down-right', 'down-left'."
    )
    is_simple_bracket: Optional[bool] = Field(
        None,
        description="True if the entire drawing is just a single grouping bracket or curly brace."
    )
    bracket_style: Optional[str] = Field(
        None,
        description="If is_simple_bracket is True, the style of bracket: 'curly', 'square', 'plain'."
    )
    bracket_side: Optional[str] = Field(
        None,
        description="If is_simple_bracket is True, which side the bracket appears relative to the text: 'left', 'right'."
    )

class DocumentLayout(BaseModel):
    page_margin_cm: float = Field(
        default=2.54,
        description="Estimated page margin in centimeters (usually between 1.5 and 3.0, default 2.54)."
    )
    line_spacing: float = Field(
        default=1.15,
        description="Line spacing: 1.0 (tight), 1.15 (normal), 1.5 (airy), 2.0 (double-spaced)."
    )
    elements: List[DocElement] = Field(
        ...,
        description="The ordered list of all document elements (text, tables, drawings, blank lines) from top to bottom of the page in reading order."
    )

# Gemini client — rebuilt on key rotation
_gemini_client = _make_gemini_client(_get_gemini_key())

UPLOAD_DIR = "uploads"
OUTPUT_DIR = "outputs"
os.makedirs(UPLOAD_DIR, exist_ok=True)
os.makedirs(OUTPUT_DIR, exist_ok=True)

app = FastAPI()
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

ALIGN_MAP = {
    "left":    WD_ALIGN_PARAGRAPH.LEFT,
    "center":  WD_ALIGN_PARAGRAPH.CENTER,
    "right":   WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}

# ── Pages ──
@app.get("/", response_class=HTMLResponse)
def landing(request: Request):
    return templates.TemplateResponse("landing.html", {"request": request})

@app.get("/convert", response_class=HTMLResponse)
def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

@app.get("/login", response_class=HTMLResponse)
def login(request: Request):
    return templates.TemplateResponse("login.html", {"request": request})

@app.get("/history", response_class=HTMLResponse)
def history(request: Request):
    return templates.TemplateResponse("history.html", {"request": request})

@app.get("/about", response_class=HTMLResponse)
def about(request: Request):
    return templates.TemplateResponse("about.html", {"request": request})


# ── PDF → Images ──
def pdf_to_images(path):
    doc = pdfium.PdfDocument(path)
    paths = []
    for i, page in enumerate(doc):
        bitmap = page.render(scale=2)
        pil_img = bitmap.to_pil()
        p = f"{UPLOAD_DIR}/{uuid.uuid4()}_{i}.png"
        pil_img.save(p)
        paths.append(p)
    return paths


# ═══════════════════════════════════════════════════════════════════════
#  COMBINED OCR PROMPT — Consolidated OCR & Layout structured detection
# ═══════════════════════════════════════════════════════════════════════
COMBINED_OCR_PROMPT = """\
You are a world-class document OCR engine and layout analyst. Your task is to extract every piece of text and structural information from the provided document image with PERFECT accuracy, preserving all content exactly as it appears.

Analyse the document strictly from TOP to BOTTOM, LEFT to RIGHT. Return elements in their exact visual reading order.

━━━ GENERAL TRANSCRIPTION RULES ━━━
1. COPY TEXT EXACTLY — do not paraphrase, summarize, correct, auto-complete, or rephrase anything. Every word, number, symbol, and abbreviation must be transcribed exactly as printed/written.
2. Preserve ALL capitalisation, punctuation, and spacing exactly as in the original.
3. If a word or character is completely illegible (cannot be determined at all), write [?].
4. NEVER invent, guess, or interpolate text that is not clearly visible in the image.
5. Preserve the EXACT spelling — even if it appears to be a typo or abbreviation.

━━━ MATHEMATICAL EQUATIONS & FORMULAS ━━━
This is CRITICAL. Mathematical content must be transcribed with perfect fidelity:
1. Transcribe all math symbols exactly: +, -, ×, ÷, ±, ≠, ≤, ≥, ∞, ∝, ∂, ∇, ∑, ∏, ∫, √, ∛
2. Greek letters: α, β, γ, δ, ε, ζ, η, θ, λ, μ, ν, π, ρ, σ, τ, φ, ψ, ω (capital: Δ, Σ, Π, Φ, Ω)
3. Superscripts: write as x^2, x^n, e^(x+1), A^T, etc.
4. Subscripts: write as x_1, a_n, C_p, log_2(n), etc.
5. Fractions: write as (numerator)/(denominator), e.g., (3x+1)/(x^2-1) or use the ÷ symbol if a simple fraction.
6. Square roots: write as √(expression), e.g., √(x^2+y^2).
7. Vectors: write with arrow notation →v or bold: **v**.
8. Absolute value: |x|, norm: ||v||.
9. Matrix notation: write row by row separated by semicolons inside square brackets: [a b; c d].
10. For multi-line equations that continue on the next line (starting with = or operator), emit each line as a SEPARATE text element with tag='BODY'.
11. NEVER render math as images — always transcribe as Unicode text.

━━━ TEXT ELEMENT TAGS ━━━
For each text element, assign the correct tag:
- HEADING: Major title or primary heading (significantly larger/bolder font than body)
- SUBHEAD: Sub-heading or section label (bold, slightly larger than body, or underlined heading)
- BODY: Normal body text, sentences, paragraphs, equations, formulas, running text
- BULLET: List item, bullet point, numbered item (e.g., starts with •, -, 1., a), etc.)
- CENTER: Text that is visually centred on the page (dates, titles, captions)
- UNDERLN: Body text that has a visible underline drawn directly beneath it

━━━ FORMATTING FIELDS ━━━
For every text element, set these fields:
- bold: true if text is printed bold or written significantly thicker
- italic: true if text is slanted/italic
- underline: true if text has an underline
- alignment: 'left' | 'center' | 'right' | 'justify'
- font_size_pt: estimated font size (16 for main heading, 13 for sub-heading, 12 for body, 10 for footnote/caption)
- left_indent_cm: estimated left indent in cm (0.0 for normal text, 0.5 for bullet/indented)

━━━ MULTI-COLUMN DOCUMENT HEADERS ━━━
If the page header has MULTIPLE content blocks side-by-side (e.g., logo on left, title in centre, date/page info on right), represent this as ONE 'table' element with borderless=true.
- Each side-by-side block = one column.
- Do NOT emit header items as separate stacked text elements.
- If a logo or stamp appears in a header cell, leave that cell text empty and emit a separate 'drawing' element for it.

━━━ TABLES ━━━
Table extraction must be PERFECT:
1. Identify ALL tables — with or without visible gridlines.
2. Extract EVERY cell's text with 100% accuracy. Do not skip, merge, or omit any cell.
3. If a table has NO visible gridlines, set borderless=true.
4. Multi-line cell text: if a cell contains text on multiple visual lines, join them with \n (e.g., "Subject Code:\n3160003").
5. Header rows: the first row is typically the header. Preserve header text exactly.
6. Empty cells: represent as empty string "".
7. Set col_alignments to the per-column alignment list (e.g., ["center","left","center"]).
8. Mathematical content in table cells: transcribe with the same math rules as above.

━━━ TIMETABLE / SCHEDULE TABLES (special rules) ━━━
For schedule or timetable tables:
- If a DATE cell (e.g., "16/06/2026") and a DAY-OF-WEEK cell (e.g., "TUESDAY") are on separate visual lines but logically belong to the same cell (spanning multiple rows), combine them into ONE cell value: "16/06/2026\nTUESDAY". Do NOT create two separate rows.
- Subject code and subject name printed on separate lines within the same cell should be combined with \n: e.g., "IPDC\n(3160003)".
- Time ranges, batch info, and lab/room info that appear in the same cell should all be included, using \n to separate visual lines.

━━━ SIGNATURE BLOCKS & SIDE-BY-SIDE COLUMNS ━━━
If the document has side-by-side text columns, key-value pairs, or signature blocks, represent them as ONE 'table' element with borderless=true. Each column = one table column. Signature/stamp drawings inside a signature block MUST be placed as a 'drawing' element within the table rows — NOT as standalone elements after the table. Leave the cell with the signature image empty; it will be matched automatically.

━━━ DRAWINGS, DIAGRAMS & SHAPES ━━━
For each drawing (diagram, flowchart, graph, chart, sketch, logo, stamp, signature):
1. Emit a 'drawing' element with a TIGHT normalized bbox [x1, y1, x2, y2] (values 0.0–1.0).
2. The bbox must be measured precisely to where the ink starts and ends. Do NOT include surrounding whitespace.
3. Do NOT use round numbers — use the actual measured fractional coordinates.
4. Simple Arrow: set is_simple_arrow=true and specify arrow_direction.
5. Simple Bracket: set is_simple_bracket=true and specify bracket_style and bracket_side.

━━━ PAGE SETTINGS ━━━
Estimate page_margin_cm (typically 1.5–3.0 cm) and line_spacing (1.0=tight, 1.15=normal, 1.5=airy, 2.0=double).
"""


PLAIN_OCR_PROMPT = (
    "Extract all the text from this handwriting image exactly as written. "
    "Read every line top to bottom. Output each line on its own line. "
    "Preserve capitalisation, punctuation, numbers and symbols exactly. "
    "Skip any diagrams or drawings silently. "
    "No descriptions, no extra formatting, no commentary."
)


# ─────────────────────────────────────────────────
#  Image preprocessing — enhance for OCR accuracy
# ─────────────────────────────────────────────────
def _preprocess_for_ocr(img: Image.Image) -> Image.Image:
    """
    Standardize the image for Gemini multimodal analysis:
    - Ensure it is in RGB format
    - Gently upscale if the image is too small (width < 1600px) to preserve detail
    - Avoid applying extreme contrast or sharpening filters to prevent vision degradation
    """
    if img.mode != "RGB":
        img = img.convert("RGB")

    min_w = 1600
    if img.width < min_w:
        scale = min_w / img.width
        new_h = int(img.height * scale)
        img = img.resize((min_w, new_h), Image.LANCZOS)
        print(f"  [PRE] Upscaled to {min_w}×{new_h}px")

    return img


# ─────────────────────────────────────────────────────────────
#  Multi-Provider Vision API — Gemini → Groq → OpenRouter
# ─────────────────────────────────────────────────────────────
_PRIMARY_MODEL  = "gemini-2.5-flash"
_FALLBACK_MODEL = "gemini-2.0-flash"

# Error patterns that mean "quota/rate-limit" → try next provider
_QUOTA_PATTERNS = (
    "429", "RESOURCE_EXHAUSTED", "quota", "rate_limit",
    "rate limit", "exceeded", "too many", "QUOTA_EXCEEDED",
)

def _is_quota_error(err: Exception) -> bool:
    s = str(err).lower()
    return any(p.lower() in s for p in _QUOTA_PATTERNS)

def _is_overload_error(err: Exception) -> bool:
    s = str(err)
    return "503" in s or "UNAVAILABLE" in s


# ── Helper: PIL Image → base64 string ──
def _img_to_base64(img: Image.Image, fmt="JPEG") -> str:
    buf = io.BytesIO()
    if img.mode != "RGB":
        img = img.convert("RGB")
    img.save(buf, format=fmt, quality=90)
    return base64.b64encode(buf.getvalue()).decode()


# ── Provider 1: Gemini (with multi-key rotation) ──
def _try_gemini(img: Image.Image, prompt, config=None, retries=4) -> str:
    global _gemini_client, _provider_status, _last_provider
    last_error = None
    models_to_try = [_PRIMARY_MODEL, _FALLBACK_MODEL]

    for model in models_to_try:
        for attempt in range(1, retries + 2):
            try:
                print(f"  [Gemini/{model}] attempt {attempt}, key idx={_gemini_key_idx % max(1,len(_GEMINI_KEYS))}...")
                resp = _gemini_client.models.generate_content(
                    model=model,
                    contents=[prompt, img],
                    config=config,
                )
                _provider_status["gemini"] = "ok"
                _last_provider = "gemini"
                return resp.text
            except Exception as e:
                last_error = e
                err_str = str(e)
                print(f"  [Gemini] attempt {attempt} failed: {type(e).__name__}: {e}")

                if _is_quota_error(e):
                    # Rotate Gemini key and retry immediately
                    if len(_GEMINI_KEYS) > 1:
                        new_key = _rotate_gemini_key()
                        _gemini_client = _make_gemini_client(new_key)
                        time.sleep(1)
                        continue  # retry same model with new key
                    else:
                        _provider_status["gemini"] = "quota_exceeded"
                        raise  # single key, escalate to next provider

                if _is_overload_error(e):
                    if attempt <= retries:
                        wait = (2 ** attempt) * 2
                        print(f"  [Gemini] overloaded, retrying in {wait}s...")
                        time.sleep(wait)
                    else:
                        _provider_status["gemini"] = "overloaded"
                        break
                else:
                    # Non-retryable error
                    _provider_status["gemini"] = "error"
                    raise

        if not _is_overload_error(last_error) if last_error else True:
            break
        if model != models_to_try[-1]:
            print(f"  [Gemini] switching to fallback model {_FALLBACK_MODEL}")

    raise last_error


# ── Provider 2: Groq (Llama 4 Scout Vision) ──
def _try_groq(img: Image.Image, prompt, config=None, retries=3) -> str:
    global _provider_status, _last_provider
    if not GROQ_API_KEY:
        raise RuntimeError("Groq API key not configured")

    try:
        from openai import OpenAI
    except ImportError:
        raise RuntimeError("openai package not installed — run: pip install openai")

    # Groq uses OpenAI-compatible API
    groq_client = OpenAI(
        api_key=GROQ_API_KEY,
        base_url="https://api.groq.com/openai/v1",
    )

    # Groq vision models to try in order
    groq_models = [
        "meta-llama/llama-4-scout-17b-16e-instruct",
        "meta-llama/llama-4-maverick-17b-128e-instruct",
        "llama-3.2-90b-vision-preview",
        "llama-3.2-11b-vision-preview",
    ]

    img_b64 = _img_to_base64(img)
    plain_prompt = (
        prompt if isinstance(prompt, str)
        else PLAIN_OCR_PROMPT
    )
    is_json = config and getattr(config, "response_mime_type", None) == "application/json"
    if is_json:
        import json
        schema_str = json.dumps(DocumentLayout.model_json_schema(), indent=2)
        plain_prompt += f"\n\nYou MUST return a JSON object matching this JSON Schema:\n{schema_str}\n\nReturn ONLY the valid JSON object, with no markdown formatting or commentary."

    last_error = None
    for model in groq_models:
        for attempt in range(1, retries + 2):
            try:
                print(f"  [Groq/{model}] attempt {attempt}...")
                kwargs = {
                    "model": model,
                    "messages": [{
                        "role": "user",
                        "content": [
                            {"type": "text", "text": plain_prompt},
                            {"type": "image_url", "image_url": {
                                "url": f"data:image/jpeg;base64,{img_b64}"
                            }},
                        ]
                    }],
                    "max_tokens": 4096,
                    "temperature": 0.1,
                }
                if is_json:
                    kwargs["response_format"] = {"type": "json_object"}
                
                try:
                    resp = groq_client.chat.completions.create(**kwargs)
                except Exception as je:
                    if is_json and "response_format" in kwargs:
                        del kwargs["response_format"]
                        resp = groq_client.chat.completions.create(**kwargs)
                    else:
                        raise je

                result = resp.choices[0].message.content
                result = _clean_json_response(result)
                _provider_status["groq"] = "ok"
                _last_provider = "groq"
                print(f"  [Groq] Success with {model}")
                return result
            except Exception as e:
                last_error = e
                print(f"  [Groq/{model}] attempt {attempt} failed: {type(e).__name__}: {e}")
                if _is_quota_error(e):
                    _provider_status["groq"] = "quota_exceeded"
                    break  # try next model
                if _is_overload_error(e) and attempt <= retries:
                    time.sleep((2 ** attempt))
                else:
                    break
    _provider_status["groq"] = "error"
    raise last_error or RuntimeError("Groq: all models failed")


# ── Provider 3: OpenRouter (free vision models) ──
def _try_openrouter(img: Image.Image, prompt, config=None, retries=3) -> str:
    global _provider_status, _last_provider
    if not OPENROUTER_API_KEY:
        raise RuntimeError("OpenRouter API key not configured")

    try:
        from openai import OpenAI
    except ImportError:
        raise RuntimeError("openai package not installed — run: pip install openai")

    or_client = OpenAI(
        api_key=OPENROUTER_API_KEY,
        base_url="https://openrouter.ai/api/v1",
        default_headers={
            "HTTP-Referer": "https://docify.app",
            "X-Title": "Docify AI",
        }
    )

    # Free vision-capable models on OpenRouter
    or_models = [
        "qwen/qwen2.5-vl-72b-instruct:free",
        "qwen/qwen2-vl-7b-instruct:free",
        "meta-llama/llama-3.2-11b-vision-instruct:free",
        "google/gemini-2.0-flash-exp:free",
    ]

    img_b64 = _img_to_base64(img)
    plain_prompt = prompt if isinstance(prompt, str) else PLAIN_OCR_PROMPT
    is_json = config and getattr(config, "response_mime_type", None) == "application/json"
    if is_json:
        import json
        schema_str = json.dumps(DocumentLayout.model_json_schema(), indent=2)
        plain_prompt += f"\n\nYou MUST return a JSON object matching this JSON Schema:\n{schema_str}\n\nReturn ONLY the valid JSON object, with no markdown formatting or commentary."

    last_error = None
    for model in or_models:
        for attempt in range(1, retries + 2):
            try:
                print(f"  [OpenRouter/{model}] attempt {attempt}...")
                kwargs = {
                    "model": model,
                    "messages": [{
                        "role": "user",
                        "content": [
                            {"type": "text", "text": plain_prompt},
                            {"type": "image_url", "image_url": {
                                "url": f"data:image/jpeg;base64,{img_b64}"
                            }},
                        ]
                    }],
                    "max_tokens": 4096,
                    "temperature": 0.1,
                }
                if is_json:
                    kwargs["response_format"] = {"type": "json_object"}
                
                try:
                    resp = or_client.chat.completions.create(**kwargs)
                except Exception as je:
                    if is_json and "response_format" in kwargs:
                        del kwargs["response_format"]
                        resp = or_client.chat.completions.create(**kwargs)
                    else:
                        raise je

                result = resp.choices[0].message.content
                result = _clean_json_response(result)
                _provider_status["openrouter"] = "ok"
                _last_provider = "openrouter"
                print(f"  [OpenRouter] Success with {model}")
                return result
            except Exception as e:
                last_error = e
                print(f"  [OpenRouter/{model}] attempt {attempt} failed: {type(e).__name__}: {e}")
                if _is_quota_error(e):
                    _provider_status["openrouter"] = "quota_exceeded"
                    break
                if _is_overload_error(e) and attempt <= retries:
                    time.sleep((2 ** attempt))
                else:
                    break
    _provider_status["openrouter"] = "error"
    raise last_error or RuntimeError("OpenRouter: all models failed")


# ── Master dispatcher: tries each provider in order ──
def _call_vision_api(img: Image.Image, prompt, config=None) -> str:
    """
    Try Gemini → Groq → OpenRouter in sequence.
    Falls to next provider only on quota/rate-limit errors.
    """
    providers = []
    if _GEMINI_KEYS:
        providers.append(("Gemini",     lambda: _try_gemini(img, prompt, config)))
    if GROQ_API_KEY:
        providers.append(("Groq",       lambda: _try_groq(img, prompt, config)))
    if OPENROUTER_API_KEY:
        providers.append(("OpenRouter", lambda: _try_openrouter(img, prompt, config)))

    if not providers:
        raise RuntimeError("No API keys configured. Add GEMINI_API_KEY, GROQ_API_KEY, or OPENROUTER_API_KEY to .env")

    last_error = None
    for name, fn in providers:
        try:
            return fn()
        except Exception as e:
            last_error = e
            if _is_quota_error(e):
                print(f"  [{name}] quota/rate-limit hit — trying next provider...")
                continue
            else:
                # Non-quota error (bad image, auth failure, etc) — re-raise immediately
                raise

    raise RuntimeError(
        f"All API providers exhausted. Last error: {last_error}"
    )


# Backward-compat alias (used by run_ocr_manual)
def _call_gemini(img, prompt, retries=4, preprocess=False, config=None):
    if preprocess:
        img = _preprocess_for_ocr(img)
    return _call_vision_api(img, prompt, config)


# ─────────────────────────────────────────────────
#  Allowed spacing values
# ─────────────────────────────────────────────────
ALLOWED_SPACING = (0, 3, 6, 9, 12, 18, 24)

def _snap_spacing(val, font_size_pt=12):
    try:
        val = float(val)
    except (TypeError, ValueError):
        val = 0.0
    return min(ALLOWED_SPACING, key=lambda x: abs(x - val))


# ─────────────────────────────────────────────────
#  OCR text post-processing cleanup
# ─────────────────────────────────────────────────
# Common Gemini OCR confusions: fix recurring misreads
_OCR_FIXES = [
    # Zero / O confusion in obvious numeric contexts
    (re.compile(r'(?<=\d)O(?=\d)'), '0'),   # 1O5 → 105
    (re.compile(r'(?<=\d)o(?=\d)'), '0'),   # 1o5 → 105
    # Trailing/leading whitespace inside parens
    (re.compile(r'\(\s+'), '('),
    (re.compile(r'\s+\)'), ')'),
    # Stray backticks from markdown leakage
    (re.compile(r'`+'), ''),
    # Remove accidental prefix repetition e.g. "BODY: BODY: text"
    (re.compile(r'^(HEADING|SUBHEAD|BODY|BULLET|CENTER|UNDERLN):\s*\1:\s*', re.I), r'\1: '),
]

def _clean_ocr_text(text: str) -> str:
    """Apply post-processing fixes to a single transcribed line."""
    for pattern, replacement in _OCR_FIXES:
        text = pattern.sub(replacement, text)
    
    text = text.rstrip()
    
    # Convert multiple interior spaces to non-breaking spaces to preserve alignment
    # e.g., "x  =  y" -> "x \xA0= \xA0y" (preserves spacing but allows word wrap on single spaces)
    while "  " in text:
        text = text.replace("  ", " \xA0")
        
    # Convert leading spaces to non-breaking spaces so DOCX/HTML don't collapse them
    match = re.match(r'^(\s+)(.*)$', text)
    if match:
        text = match.group(1).replace(" ", "\xA0") + match.group(2)
        
    return text


# ─────────────────────────────────────────────────
#  Parse raw OCR text lines into structured elements
# ─────────────────────────────────────────────────
# Map prefix → (bold, italic, underline, font_size_pt, alignment, left_indent_cm, space_before_pt, space_after_pt)
_PREFIX_MAP = {
    "HEADING":  (True,  False, False, 16, "left",   0.0, 12, 6),
    "SUBHEAD":  (True,  False, False, 13, "left",   0.0,  8, 4),
    "BODY":     (False, False, False, 12, "left",   0.0,  0, 3),
    "BULLET":   (False, False, False, 12, "left",   0.5,  0, 3),
    "CENTER":   (False, False, False, 12, "center", 0.0,  4, 4),
    "UNDERLN":  (False, False, True,  12, "left",   0.0,  0, 3),
    "TABLE":    (False, False, False, 11, "left",   0.0,  4, 4),
}

def _parse_text_lines(raw_text):
    """
    Convert the structured OCR output into a list of element dicts.
    Handles all PREFIX: tags produced by TEXT_OCR_PROMPT.
    Falls back gracefully for un-tagged lines (treated as BODY).
    """
    elements = []
    # Collapse runs of 3+ blank lines into 2 maximum
    collapsed = re.sub(r'\n{3,}', '\n\n', raw_text)

    last_equal_idx = -1

    for line in collapsed.split("\n"):
        line_clean = line.rstrip()

        # True blank line → paragraph gap
        if line_clean.strip() == "":
            last_equal_idx = -1
            # Avoid stacking multiple blank_lines
            if elements and elements[-1].get("type") != "blank_line":
                elements.append({"type": "blank_line"})
            continue

        # Detect and strip prefix smoothly while preserving leading spaces
        prefix = "BODY"
        text = line_clean
        for pfx in _PREFIX_MAP:
            # Match prefix (case-insensitive), an optional colon-space, and capture the rest
            pattern = rf'^\s*{pfx}:\s?(.*)$'
            match = re.match(pattern, line_clean, re.IGNORECASE)
            if match:
                prefix = pfx
                text = match.group(1)  # Preserves leading spaces of the actual content
                break

        # Clean the transcribed text
        text = _clean_ocr_text(text)

        if not text:
            continue

        if prefix == "TABLE":
            last_equal_idx = -1
            # Parse markdown row: | col1 | col2 |
            # Clean up the \xA0 spaces specifically for table cells
            clean_row = text.replace('\xA0', ' ').strip()
            cells = [c.strip() for c in clean_row.strip('|').split('|')]
            
            # Check if it's a markdown separator row (e.g., |---|---| )
            is_separator = all(re.match(r'^[-:]+$', c) for c in cells if c)
            
            if not is_separator:
                if elements and elements[-1].get("type") == "table":
                    elements[-1]["data"].append(cells)
                else:
                    elements.append({
                        "type": "table",
                        "data": [cells]
                    })
            continue

        # ── Smart Math Equation '=' Alignment ──
        # If the line starts with '=' (ignoring spaces), pad it so the '='
        # aligns horizontally with the '=' in the previous equation line.
        temp_text = text.replace("\xA0", " ")
        first_char_idx = len(temp_text) - len(temp_text.lstrip())
        
        if first_char_idx < len(temp_text) and temp_text[first_char_idx] == '=' and last_equal_idx > 0:
            # It's a continuation line starting with =
            needed_spaces = last_equal_idx - first_char_idx
            if needed_spaces > 0:
                text = ("\xA0" * needed_spaces) + text
        else:
            # Look for an '=' in a normal line to anchor future continuations
            eq_pos = temp_text.find("=")
            if eq_pos > 0:
                last_equal_idx = eq_pos

        bold, italic, underline, fsize, align, left_i, sp_bef, sp_aft = _PREFIX_MAP[prefix]

        # Add space_before for headings only if previous element was not blank
        if prefix in ("HEADING", "SUBHEAD") and elements and elements[-1].get("type") != "blank_line":
            sp_bef = max(sp_bef, 10)

        elements.append({
            "type":                 "text",
            "text":                 text,
            "alignment":            align,
            "font_size_pt":         fsize,
            "bold":                 bold,
            "italic":               italic,
            "underline":            underline,
            "first_line_indent_cm": 0.0,
            "left_indent_cm":       left_i,
            "space_before_pt":      sp_bef,
            "space_after_pt":       sp_aft,
        })
    return elements


# ─────────────────────────────────────────────────
#  Arrow Unicode mapping
# ─────────────────────────────────────────────────
_ARROW_UNICODE = {
    "right":      "⟶",   # Long rightwards arrow
    "left":       "⟵",   # Long leftwards arrow
    "up":         "↑",    # Upwards arrow
    "down":       "↓",    # Downwards arrow
    "up-right":   "↗",   # North-east arrow
    "up-left":    "↖",   # North-west arrow
    "down-right": "↘",   # South-east arrow
    "down-left":  "↙",   # South-west arrow
}

def _resolve_arrow_char(direction: str) -> str:
    """Return the best Unicode arrow for the given direction string."""
    d = (direction or "").strip().lower()
    return _ARROW_UNICODE.get(d, "⟶")   # default: rightwards


# ─────────────────────────────────────────────────
#  Bracket / Brace Unicode mapping
# ─────────────────────────────────────────────────
# Maps (style, side) -> (display_char, font_size_multiplier)
_BRACKET_CHARS = {
    # curly braces
    ("curly",  "left"):  "⎧⎨⎩",   # top, mid, bottom curly bracket pieces
    ("curly",  "right"): "⎫⎬⎭",
    # square brackets
    ("square", "left"):  "[",
    ("square", "right"): "]",
    # plain / L-bracket
    ("plain",  "left"):  "⎡⎣",   # top and bottom square bracket pieces
    ("plain",  "right"): "⎤⎦",
}

# Single display chars for simple cases (used for PDF/DOCX single-char rendering)
_BRACKET_SINGLE = {
    ("curly",  "left"):  "{",
    ("curly",  "right"): "}",
    ("square", "left"):  "[",
    ("square", "right"): "]",
    ("plain",  "left"):  "[",
    ("plain",  "right"): "]",
}

def _resolve_bracket_char(style: str, side: str) -> str:
    """Return the best single bracket character for style+side."""
    s = (style or "square").strip().lower()
    d = (side  or "left").strip().lower()
    return _BRACKET_SINGLE.get((s, d), "[")


# ─────────────────────────────────────────────────
#  Merge text elements with drawing inserts
#  Drawings are inserted at the correct Y position
#  in the text flow using position_y_fraction.
# ─────────────────────────────────────────────────
def _merge_elements(text_elements, drawings):
    """
    Insert drawing elements into the text flow at the position that matches
    their vertical location in the source image.

    Simple-arrow drawings become type="arrow" (rendered as Unicode text).
    Real drawings become type="drawing" (rendered as inline cropped image).
    """
    if not drawings:
        return text_elements

    text_only = [e for e in text_elements if e.get("type") == "text"]
    total = max(1, len(text_only))

    insertions = []
    for d in sorted(drawings, key=lambda x: x.get("position_y_fraction", 0.5)):
        frac = float(d.get("position_y_fraction", 0.5))
        idx  = int(frac * total)
        if d.get("is_simple_arrow", False):
            # Represent as a Unicode arrow text element
            arrow_char = _resolve_arrow_char(d.get("arrow_direction", "right"))
            insertions.append((idx, {
                "type":        "arrow",
                "arrow_char":  arrow_char,
                "description": d.get("description", ""),
            }))
            print(f"  [MERGE] Arrow-only drawing -> Unicode '{arrow_char}'")
        elif d.get("is_simple_bracket", False):
            # Represent as a Unicode bracket character
            brk_char = _resolve_bracket_char(
                d.get("bracket_style", "square"),
                d.get("bracket_side",  "left"),
            )
            insertions.append((idx, {
                "type":          "bracket",
                "bracket_char":  brk_char,
                "bracket_style": d.get("bracket_style", "square"),
                "bracket_side":  d.get("bracket_side",  "left"),
                "description":   d.get("description", ""),
            }))
            print(f"  [MERGE] Bracket-only drawing -> Unicode '{brk_char}'")
        else:
            insertions.append((idx, {
                "type":        "drawing",
                "bbox":        d.get("bbox"),
                "description": d.get("description", ""),
            }))

    # Walk through text_elements, mapping their text-only index to a position
    result = []
    text_idx = 0          # counts only "text" type elements seen so far
    ins_ptr  = 0          # pointer into insertions list

    for el in text_elements:
        # Before adding this element, check if any drawing should be inserted
        while ins_ptr < len(insertions) and insertions[ins_ptr][0] <= text_idx:
            _, drw = insertions[ins_ptr]
            result.append({
                "type": "drawing",
                "bbox": drw.get("bbox"),
                "description": drw.get("description", ""),
            })
            ins_ptr += 1

        result.append(el)
        if el.get("type") == "text":
            text_idx += 1

    # Append any remaining drawings at the end
    while ins_ptr < len(insertions):
        _, drw = insertions[ins_ptr]
        result.append({
            "type": "drawing",
            "bbox": drw.get("bbox"),
            "description": drw.get("description", ""),
        })
        ins_ptr += 1

    return result


# ─────────────────────────────────────────────────
#  Auto OCR — two-step pipeline
#  Step 1: pure text OCR (high accuracy)
#  Step 2: layout/drawing detection
#  Then merge.
# ─────────────────────────────────────────────────
def run_ocr_auto(image_path):
    import hashlib
    import json
    import os

    # Simple hash-based local cache for OCR
    try:
        hasher = hashlib.md5()
        with open(image_path, "rb") as f:
            hasher.update(f.read())
        cache_key = hasher.hexdigest()
        cache_file = os.path.join(UPLOAD_DIR, f"ocr_cache_{cache_key}.json")
        if os.path.exists(cache_file):
            print(f"  [CACHE] Loading cached OCR layout from {cache_file}...")
            with open(cache_file, "r", encoding="utf-8") as f:
                cached_data = json.load(f)
                return cached_data["elements"], cached_data["meta"]
    except Exception as e:
        print(f"  [CACHE] Failed to check/load cache: {e}")

    print(f"[AUTO] Consolidated OCR & Layout: {image_path}")
    raw_img = Image.open(image_path)

    # Standardize image (convert to RGB, upscale if too small)
    img = _preprocess_for_ocr(raw_img)

    try:
        print("  Calling vision API for combined layout and text OCR...")
        # Structured JSON config — only Gemini supports native schema;
        # Groq/OpenRouter will fall back to plain text pipeline automatically.
        config = None
        if _GEMINI_KEYS and _provider_status.get("gemini") not in ("quota_exceeded",):
            config = genai_types.GenerateContentConfig(
                response_mime_type="application/json",
                response_schema=DocumentLayout,
                temperature=0.0,  # Maximum determinism for OCR accuracy
            )

        resp_text = _call_vision_api(img, COMBINED_OCR_PROMPT, config=config)
        resp_text = _clean_json_response(resp_text)
        layout_data = DocumentLayout.model_validate_json(resp_text)
        print(f"  Gemini Layout: margin={layout_data.page_margin_cm}cm, spacing={layout_data.line_spacing}")
        
        elements = []
        for el in layout_data.elements:
            etype = el.type
            if etype == "blank_line":
                elements.append({"type": "blank_line"})
            elif etype == "table":
                elements.append({
                    "type": "table",
                    "data": el.table_data or [],
                    "borderless": getattr(el, "borderless", False) or False,
                    "col_alignments": getattr(el, "col_alignments", None) or [],
                })
            elif etype == "drawing":
                if el.is_simple_arrow:
                    arrow_char = _resolve_arrow_char(el.arrow_direction or "right")
                    elements.append({
                        "type":        "arrow",
                        "arrow_char":  arrow_char,
                        "description": el.description or "",
                    })
                elif el.is_simple_bracket:
                    brk_char = _resolve_bracket_char(
                        el.bracket_style or "square",
                        el.bracket_side or "left"
                    )
                    elements.append({
                        "type":          "bracket",
                        "bracket_char":  brk_char,
                        "bracket_style": el.bracket_style or "square",
                        "bracket_side":  el.bracket_side or "left",
                        "description":   el.description or "",
                    })
                else:
                    elements.append({
                        "type":        "drawing",
                        "bbox":        el.bbox,
                        "description": el.description or "",
                    })
            elif etype == "text":
                tag = (el.tag or "BODY").upper()
                text = _clean_ocr_text(el.text or "")
                if not text:
                    continue
                
                # Fetch settings from map
                bold_mapped, italic_mapped, underline_mapped, fsize, align, left_i, sp_bef, sp_aft = _PREFIX_MAP.get(
                    tag, _PREFIX_MAP["BODY"]
                )
                
                # Use model-detected override values if specified, else fall back to prefix defaults
                bold = el.bold if el.bold is not None else bold_mapped
                italic = el.italic if el.italic is not None else italic_mapped
                underline = el.underline if el.underline is not None else underline_mapped
                align = el.alignment if el.alignment is not None else align
                left_i = el.left_indent_cm if el.left_indent_cm is not None else left_i
                # Use model-reported font size if available and plausible
                if el.font_size_pt is not None and 6.0 <= el.font_size_pt <= 36.0:
                    fsize = float(el.font_size_pt)
                
                elements.append({
                    "type":                 "text",
                    "text":                 text,
                    "alignment":            align,
                    "font_size_pt":         fsize,
                    "bold":                 bold,
                    "italic":               italic,
                    "underline":            underline,
                    "first_line_indent_cm": 0.0,
                    "left_indent_cm":       left_i,
                    "space_before_pt":      sp_bef,
                    "space_after_pt":       sp_aft,
                })
        
        meta = {
            "line_spacing": layout_data.line_spacing,
            "page_margin_cm": layout_data.page_margin_cm
        }
        try:
            with open(cache_file, "w", encoding="utf-8") as f:
                json.dump({"elements": elements, "meta": meta}, f, ensure_ascii=False, indent=2)
            print(f"  [CACHE] Saved OCR layout cache to {cache_file}")
        except Exception as ce:
            print(f"  [CACHE] Failed to save cache: {ce}")
        return elements, meta

    except Exception as e:
        print(f"  Combined OCR failed ({e}), trying plain OCR fallback...")
        try:
            raw_text = _call_gemini(img, PLAIN_OCR_PROMPT)
            text_elements = _parse_text_lines(raw_text)
            meta = {"line_spacing": 1.15, "page_margin_cm": 2.54}
            try:
                with open(cache_file, "w", encoding="utf-8") as f:
                    json.dump({"elements": text_elements, "meta": meta}, f, ensure_ascii=False, indent=2)
                print(f"  [CACHE] Saved OCR fallback cache to {cache_file}")
            except Exception as ce:
                print(f"  [CACHE] Failed to save fallback cache: {ce}")
            return text_elements, meta
        except Exception as e2:
            err = f"[OCR FAILED: {type(e2).__name__}: {str(e2)[:200]}]"
            return [{"type": "text", "text": err, "alignment": "left",
                     "font_size_pt": 12, "bold": False, "italic": False,
                     "underline": False, "first_line_indent_cm": 0.0,
                     "left_indent_cm": 0.0, "space_before_pt": 0, "space_after_pt": 0}], \
                   {"line_spacing": 1.15, "page_margin_cm": 2.54}


# ─────────────────────────────────────────────────
#  Manual OCR — plain text only
# ─────────────────────────────────────────────────
def run_ocr_manual(image_path):
    print(f"[MANUAL] Plain OCR: {image_path}")
    img = Image.open(image_path)
    try:
        raw = _call_gemini(img, PLAIN_OCR_PROMPT)
        return [l.strip() for l in raw.split("\n") if l.strip()]
    except Exception as e:
        return [f"[OCR FAILED: {type(e).__name__}: {str(e)[:200]}]"]


# ─────────────────────────────────────────────────
#  Drawing crop helper
# ─────────────────────────────────────────────────
def _crop_drawing(image_path, bbox, padding=0.05, pad_px=0,
                  pad_top=None, pad_bottom=None, pad_left=None, pad_right=None):
    """
    Crop the drawing region from image_path using normalised bbox [x1,y1,x2,y2].
    padding: relative fraction to expand each side (default 5% so drawing edges are never clipped).
    pad_px: absolute pixel padding on all sides (overridden per-side by pad_top/bottom/left/right).
    pad_top/bottom/left/right: per-side absolute pixel padding overrides.
    Returns the path to a cropped temp PNG, or None on any failure.
    """
    try:
        if not bbox or len(bbox) != 4:
            return None
        x1n, y1n, x2n, y2n = [float(v) for v in bbox]
        img = Image.open(image_path)
        w, h = img.size
        if padding > 0.0:
            bw = x2n - x1n
            bh = y2n - y1n
            x1n -= bw * padding
            y1n -= bh * padding
            x2n += bw * padding
            y2n += bh * padding
        # Absolute pixel padding with per-side overrides
        pt = (pad_top    if pad_top    is not None else pad_px)
        pb = (pad_bottom if pad_bottom is not None else pad_px)
        pl = (pad_left   if pad_left   is not None else pad_px)
        pr = (pad_right  if pad_right  is not None else pad_px)
        x1n -= pl / w
        y1n -= pt / h
        x2n += pr / w
        y2n += pb / h
        x1n = max(0.0, min(1.0, x1n))
        y1n = max(0.0, min(1.0, y1n))
        x2n = max(0.0, min(1.0, x2n))
        y2n = max(0.0, min(1.0, y2n))
        if x2n <= x1n or y2n <= y1n:
            return None
        cropped = img.crop((int(x1n*w), int(y1n*h), int(x2n*w), int(y2n*h)))
        out_path = f"{UPLOAD_DIR}/{uuid.uuid4()}_crop.png"
        cropped.save(out_path)
        print(f"  Cropped drawing: {out_path} ({int((x2n-x1n)*w)}\u00d7{int((y2n-y1n)*h)}px)")
        return out_path
    except Exception as e:
        print(f"  _crop_drawing failed: {e}")
        return None

def _sanitize_xml_text(val: str) -> str:
    if not val:
        return ""
    # Filter out control characters that are invalid in XML:
    # Valid characters: 0x09 (tab), 0x0A (LF), 0x0D (CR), 0x20 to 0xD7FF, 0xE000 to 0xFFFD, 0x10000 to 0x10FFFF
    return "".join(
        c for c in val
        if ord(c) in (0x09, 0x0A, 0x0D) or (0x20 <= ord(c) <= 0xD7FF) or (0xE000 <= ord(c) <= 0xFFFD) or (0x10000 <= ord(c) <= 0x10FFFF)
    )

def _sanitize_elements(elements):
    """Sanitize all text fields in all layout elements to ensure XML compatibility."""
    if not elements:
        return
    for el in elements:
        # Sanitize 'text' field
        if "text" in el and el["text"]:
            el["text"] = _sanitize_xml_text(el["text"])
        
        # Sanitize 'table_data' field
        if "data" in el and el["data"]:
            sanitized_data = []
            for row in el["data"]:
                sanitized_row = []
                for cell in row:
                    sanitized_row.append(_sanitize_xml_text(cell))
                sanitized_data.append(sanitized_row)
            el["data"] = sanitized_data
            
        # Sanitize other potential string fields
        for key in ["description", "bracket_char", "arrow_char"]:
            if key in el and el[key]:
                el[key] = _sanitize_xml_text(el[key])


# ─────────────────────────────────────────────────
#  Docx helpers
# ─────────────────────────────────────────────────
def _apply_run(run, font_family, font_size_pt, bold=False, italic=False, underline=False):
    run.font.name      = font_family
    run.font.size      = Pt(font_size_pt)
    run.font.bold      = bold
    run.font.italic    = italic
    run.font.underline = underline


def _apply_para(para, alignment_str, line_spacing,
                space_before_pt=0, space_after_pt=0,
                first_line_indent_cm=0.0, left_indent_cm=0.0):
    fmt = para.paragraph_format
    fmt.alignment         = ALIGN_MAP.get(alignment_str, WD_ALIGN_PARAGRAPH.LEFT)
    fmt.line_spacing      = line_spacing
    fmt.space_before      = Pt(space_before_pt)
    fmt.space_after       = Pt(space_after_pt)
    fmt.first_line_indent = Cm(first_line_indent_cm)
    fmt.left_indent       = Cm(left_indent_cm)


def _set_margins(doc, margin_cm):
    margin_cm = max(1.0, min(5.0, float(margin_cm)))
    for section in doc.sections:
        section.top_margin    = Cm(margin_cm)
        section.bottom_margin = Cm(margin_cm)
        section.left_margin   = Cm(margin_cm)
        section.right_margin  = Cm(margin_cm)


def _set_row_height(row, height_cm):
    """Set a minimum row height on a Word table row using XML trHeight."""
    tr = row._tr
    trPr = tr.find(qn('w:trPr'))
    if trPr is None:
        trPr = OxmlElement('w:trPr')
        tr.insert(0, trPr)
    trHeight = OxmlElement('w:trHeight')
    # height in twentieths of a point (1 cm = 567 twips)
    twips = int(height_cm * 567)
    trHeight.set(qn('w:val'), str(twips))
    trHeight.set(qn('w:hRule'), 'atLeast')  # minimum height, allows expansion
    trPr.append(trHeight)


def _set_cell_margins(cell, top_pt=6, bottom_pt=6, left_pt=4, right_pt=4):
    """Set cell padding (top/bottom/left/right) in points on a Word table cell."""
    tc = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tcMar = OxmlElement('w:tcMar')
    # 1 pt = 20 twips
    for side, val_pt in [('top', top_pt), ('bottom', bottom_pt), ('left', left_pt), ('right', right_pt)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(int(val_pt * 20)))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    # Remove existing tcMar if any
    existing = tcPr.find(qn('w:tcMar'))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(tcMar)

def _add_inline_picture_docx(doc, img_path, bbox, page_margin_cm=2.54):
    """
    Add a drawing INLINE in the text flow (not floating).
    Size is derived from the actual cropped pixel dimensions (not raw bbox fractions),
    so large/small bboxes don't distort the drawing size.
    Width defaults to 60% of the content area, capped at 100%.
    Aspect ratio is always preserved.
    """
    cropped = _crop_drawing(img_path, bbox)
    if not cropped:
        return

    try:
        content_w_cm = 21.0 - 2.0 * page_margin_cm
        max_w_cm = max(2.0, content_w_cm)

        # Use actual pixel dimensions to derive aspect ratio — never trust raw bbox size
        with Image.open(cropped) as ci:
            nat_w, nat_h = ci.size
        aspect = nat_h / nat_w if nat_w > 0 else 1.0

        # Size derived from relative bbox width fraction of content area
        x1n, y1n, x2n, y2n = [float(v) for v in bbox]
        bbox_w_frac = x2n - x1n
        img_w_cm = content_w_cm * bbox_w_frac
        img_w_cm = max(1.5, min(max_w_cm, img_w_cm))
        img_h_cm = img_w_cm * aspect

        # Clamp height so very tall drawings don't overflow the page
        max_h_cm = 29.7 - 2.0 * page_margin_cm
        if img_h_cm > max_h_cm * 0.65:
            img_h_cm = max_h_cm * 0.65
            img_w_cm = img_h_cm / aspect if aspect > 0 else img_w_cm

        # Horizontal alignment: mirror the original horizontal position in the source
        cx_norm = (x1n + x2n) / 2.0
        if cx_norm < 0.35:
            align = WD_ALIGN_PARAGRAPH.LEFT
        elif cx_norm > 0.65:
            align = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            align = WD_ALIGN_PARAGRAPH.CENTER

        para = doc.add_paragraph()
        para.paragraph_format.space_before = Pt(8)
        para.paragraph_format.space_after  = Pt(8)
        para.paragraph_format.alignment    = align
        run = para.add_run()
        run.add_picture(cropped, width=Cm(img_w_cm), height=Cm(img_h_cm))
        print(f"  [DOCX] Drawing inline: {img_w_cm:.2f}×{img_h_cm:.2f}cm (aspect={aspect:.2f})")
    except Exception as e:
        print(f"  [DOCX] Inline image failed: {e}")


# ─────────────────────────────────────────────────
#  PDF helpers
# ─────────────────────────────────────────────────
PT_PER_CM = 28.3465

# ── Register Unicode TTF fonts for PDF (so math/Greek/special chars render) ──
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

_PDF_FONT_NORMAL = "Helvetica"       # fallback
_PDF_FONT_BOLD   = "Helvetica-Bold"
_PDF_FONT_ITALIC = "Helvetica-Oblique"
_PDF_FONT_BOLDITALIC = "Helvetica-BoldOblique"

def _register_pdf_fonts():
    """Try to register Calibri (same as DOCX) or Arial as Unicode TTF fonts."""
    global _PDF_FONT_NORMAL, _PDF_FONT_BOLD, _PDF_FONT_ITALIC, _PDF_FONT_BOLDITALIC
    import os
    font_dir = r"C:\Windows\Fonts"
    candidates = [
        # (normal, bold, italic, bolditalic, name_prefix)
        ("calibri.ttf", "calibrib.ttf", "calibrii.ttf", "calibriz.ttf", "Calibri"),
        ("arial.ttf",   "arialbd.ttf",  "ariali.ttf",   "arialbi.ttf",  "Arial"),
    ]
    for norm, bold, ital, boldital, prefix in candidates:
        np = os.path.join(font_dir, norm)
        bp = os.path.join(font_dir, bold)
        ip = os.path.join(font_dir, ital)
        bip = os.path.join(font_dir, boldital)
        if os.path.exists(np):
            try:
                pdfmetrics.registerFont(TTFont(f"{prefix}",         np))
                pdfmetrics.registerFont(TTFont(f"{prefix}-Bold",    bp  if os.path.exists(bp)  else np))
                pdfmetrics.registerFont(TTFont(f"{prefix}-Italic",  ip  if os.path.exists(ip)  else np))
                pdfmetrics.registerFont(TTFont(f"{prefix}-BoldItalic", bip if os.path.exists(bip) else (bp if os.path.exists(bp) else np)))
                from reportlab.pdfbase.pdfmetrics import registerFontFamily
                registerFontFamily(prefix,
                    normal=prefix,
                    bold=f"{prefix}-Bold",
                    italic=f"{prefix}-Italic",
                    boldItalic=f"{prefix}-BoldItalic")
                _PDF_FONT_NORMAL     = prefix
                _PDF_FONT_BOLD       = f"{prefix}-Bold"
                _PDF_FONT_ITALIC     = f"{prefix}-Italic"
                _PDF_FONT_BOLDITALIC = f"{prefix}-BoldItalic"
                print(f"  [PDF-FONT] Registered {prefix} as Unicode PDF font")
                return
            except Exception as e:
                print(f"  [PDF-FONT] Failed to register {prefix}: {e}")
    print("  [PDF-FONT] Using Helvetica fallback (limited Unicode)")

_register_pdf_fonts()


def _rl_font(bold, italic):
    if bold and italic: return _PDF_FONT_BOLDITALIC
    if bold:            return _PDF_FONT_BOLD
    if italic:          return _PDF_FONT_ITALIC
    return _PDF_FONT_NORMAL


def _pdf_draw_text(c, txt, x, y, font_name, font_size, underline=False):
    c.setFont(font_name, font_size)
    c.drawString(x, y, txt)
    if underline:
        tw = c.stringWidth(txt, font_name, font_size)
        c.setLineWidth(0.5)
        c.line(x, y - 1.5, x + tw, y - 1.5)


def _pdf_draw_table(c, table_data, y, margin_pts, page_w, page_h, borderless=False, drawings=None, pg_path=None, table_el=None):
    """
    Draw a crisp native table on the PDF canvas with text wrapping, vertical cell merging (rowspan),
    and horizontal cell merging (colspan).
    Returns the new y position.
    """
    if not table_data:
        return y

    rows = len(table_data)
    cols = max((len(r) for r in table_data), default=0)
    if cols == 0:
        return y

    # Make a copy of table_data so we can safely update text
    table_data = [list(r) for r in table_data]
    
    # Pre-process DATE column (index 0) to combine date and day of week
    # Handles both old split-row format AND new \n-separator format from Gemini.
    DAYS_OF_WEEK = {"monday", "tuesday", "wednesday", "thursday", "friday", "saturday", "sunday"}
    for r_idx in range(1, rows):
        if len(table_data[r_idx]) == 0:
            continue
        val_curr = table_data[r_idx][0].strip()
        # Case 1: Gemini already combined with \n — normalize the separator
        if '\n' in val_curr:
            # Already combined; replace \n with space for clean display
            table_data[r_idx][0] = val_curr.replace('\n', ' ')
            continue
        # Case 2: Old format — day-of-week is on its own row
        if val_curr.lower() in DAYS_OF_WEEK:
            # Find the first non-empty cell above it in DATE column
            prev_r = r_idx - 1
            while prev_r >= 0 and (len(table_data[prev_r]) == 0 or table_data[prev_r][0].strip() == ""):
                prev_r -= 1
            if prev_r >= 0:
                val_prev = table_data[prev_r][0].strip()
                if val_curr.lower() not in val_prev.lower():
                    table_data[prev_r][0] = (val_prev + " " + val_curr).strip()
                    table_data[r_idx][0] = ""

    # Match drawings to empty cells using pre-scan assignments
    cell_drawings = {}
    if drawings and table_el:
        for d in drawings:
            if d.get("matched_table_id") == id(table_el):
                cell_drawings[d.get("matched_cell")] = d
                d["consumed"] = True

    # 1a. Detect vertical cell spans (rowspans)
    spans = [[(1, 1) for _ in range(cols)] for _ in range(rows)]
    covered = [[False for _ in range(cols)] for _ in range(rows)]
    ref_table = [list(r) for r in table_data]
    
    for col_idx in range(cols):
        # Find non-empty rows for this column (excluding header)
        non_empty_rows = []
        for r_idx in range(1, rows):
            cell_val = ref_table[r_idx][col_idx] if col_idx < len(ref_table[r_idx]) else ""
            if cell_val.strip() != "":
                non_empty_rows.append(r_idx)
                
        if not non_empty_rows:
            continue
            
        assignments = {}
        for r_idx in range(1, rows):
            cell_val = ref_table[r_idx][col_idx] if col_idx < len(ref_table[r_idx]) else ""
            if r_idx in non_empty_rows:
                assignments[r_idx] = r_idx
            else:
                # Find nearest above and below in non_empty_rows
                r_above = None
                for rx in reversed(non_empty_rows):
                    if rx < r_idx:
                        r_above = rx
                        break
                r_below = None
                for rx in non_empty_rows:
                    if rx > r_idx:
                        r_below = rx
                        break
                        
                if r_above is not None and r_below is not None:
                    # Calculate similarity score based on compatibility with other columns
                    score_above = 0
                    score_below = 0
                    for cx in range(cols):
                        if cx == col_idx:
                            continue
                        
                        val_r = ref_table[r_idx][cx].strip().lower() if cx < len(ref_table[r_idx]) else ""
                        val_above = ref_table[r_above][cx].strip().lower() if cx < len(ref_table[r_above]) else ""
                        val_below = ref_table[r_below][cx].strip().lower() if cx < len(ref_table[r_below]) else ""
                        
                        if val_r != "":
                            if val_r == val_above:
                                score_above += 1
                            if val_r == val_below:
                                score_below += 1
                    
                    if score_below > score_above:
                        assignments[r_idx] = r_below
                    else:
                        assignments[r_idx] = r_above
                elif r_above is not None:
                    assignments[r_idx] = r_above
                elif r_below is not None:
                    assignments[r_idx] = r_below
                    
        # Group by parent row to find spans
        groups = {}
        for r_idx, parent in assignments.items():
            groups.setdefault(parent, []).append(r_idx)
            
        for parent, row_list in groups.items():
            row_list.sort()
            start_r = row_list[0]
            span = len(row_list)
            
            spans[start_r][col_idx] = (span, 1)
            parent_text = ref_table[parent][col_idx] if col_idx < len(ref_table[parent]) else ""
            
            # Guard: row may be shorter than col_idx in a jagged OCR table
            if col_idx < len(table_data[start_r]):
                table_data[start_r][col_idx] = parent_text
            for rx in row_list:
                if rx != start_r:
                    covered[rx][col_idx] = True
                    if col_idx < len(table_data[rx]):
                        table_data[rx][col_idx] = ""

    # 1b. Detect horizontal cell spans (colspans) on top of vertical spans
    for r_idx in range(rows):
        col_idx = 0
        while col_idx < cols:
            if not covered[r_idx][col_idx]:
                r_span, c_span = spans[r_idx][col_idx]
                next_c = col_idx + 1
                while next_c < cols:
                    # Merge next_c if it is empty, not covered, not matched to a drawing, and has same r_span
                    # Guard: row may be shorter than next_c in a jagged OCR table
                    next_c_val = table_data[r_idx][next_c] if next_c < len(table_data[r_idx]) else None
                    if (next_c_val is not None and
                        not covered[r_idx][next_c] and 
                        (next_c_val.strip() == "") and 
                        (r_idx, next_c) not in cell_drawings and
                        spans[r_idx][next_c][0] == r_span):
                        
                        # Merge next_c into col_idx for all rows in the vertical span
                        for rx in range(r_idx, r_idx + r_span):
                            covered[rx][next_c] = True
                            if next_c < len(table_data[rx]):
                                table_data[rx][next_c] = ""
                        
                        c_span += 1
                        spans[r_idx][col_idx] = (r_span, c_span)
                        next_c += 1
                    else:
                        break
                col_idx += c_span
            else:
                col_idx += 1

    # 2. Width & Padding configs
    pad_x = 4.0 if rows > 5 else 5.0
    pad_y = 8.0 if rows > 5 else 10.0   # generous top/bottom padding like notebook rows
    fsize = 9.5 if rows > 5 else 11.0   # match DOCX cell font sizes
    font_name = _PDF_FONT_NORMAL
    font_bold = _PDF_FONT_BOLD
    
    # Calculate initial column widths, distributing spanned widths
    col_widths = [0.0] * cols
    for r_idx, r in enumerate(table_data):
        f = font_bold if r_idx == 0 else font_name
        for i, cell in enumerate(r):
            if i < cols and not covered[r_idx][i]:
                r_span, c_span = spans[r_idx][i]
                w = c.stringWidth(cell, f, fsize) + pad_x * 2
                if c_span == 1:
                    col_widths[i] = max(col_widths[i], w)
                else:
                    w_per_col = w / c_span
                    for cx in range(i, i + c_span):
                        col_widths[cx] = max(col_widths[cx], w_per_col)

    # Ensure min width for all columns
    for i in range(cols):
        col_widths[i] = max(col_widths[i], 15.0)

    usable_w = page_w - 2.0 * margin_pts
    total_w = sum(col_widths)
    
    # Scale or stretch widths to fit page logically
    if total_w < usable_w * 0.5:
        scale = (usable_w * 0.5) / total_w
        col_widths = [cw * scale for cw in col_widths]
    elif total_w > usable_w:
        scale = usable_w / total_w
        col_widths = [cw * scale for cw in col_widths]
        
    x_start = margin_pts

    # Helper to wrap text into lines, splitting by newline first
    def wrap_text(text, width, font, size):
        if not text:
            return []
        
        # Split by explicit newlines first, then wrap each line segment
        lines = []
        raw_lines = str(text).split("\n")
        for raw_line in raw_lines:
            words = raw_line.split(" ")
            cur_line = ""
            for word in words:
                test = (cur_line + " " + word).strip()
                if c.stringWidth(test, font, size) <= width - pad_x * 2:
                    cur_line = test
                else:
                    if cur_line:
                        lines.append(cur_line)
                    cur_line = word
            if cur_line:
                lines.append(cur_line)
        return lines

    # Pre-calculate wrapped lines for all active cells using spanned widths
    wrapped_data = []
    for r_idx, r in enumerate(table_data):
        row_wrapped = []
        f = font_bold if r_idx == 0 else font_name
        for col_idx in range(cols):
            if not covered[r_idx][col_idx]:
                r_span, c_span = spans[r_idx][col_idx]
                cell_txt = r[col_idx] if col_idx < len(r) else ""
                cw_span = sum(col_widths[col_idx : col_idx + c_span])
                lines = wrap_text(cell_txt, cw_span, f, fsize)
                row_wrapped.append(lines)
            else:
                row_wrapped.append([])
        wrapped_data.append(row_wrapped)

    # Calculate row heights dynamically, distributing spans correctly
    MIN_ROW_H = 28.0   # ~0.99cm minimum — matches notebook line spacing and DOCX 0.85cm rows
    row_heights = [max(MIN_ROW_H, fsize + pad_y * 2)] * rows
    line_h = fsize * 1.2 if rows > 5 else fsize * 1.3
    
    for r_idx in range(rows):
        for col_idx in range(cols):
            if not covered[r_idx][col_idx]:
                r_span, c_span = spans[r_idx][col_idx]
                max_lines = len(wrapped_data[r_idx][col_idx])
                if max_lines == 0:
                    max_lines = 1
                req_h = max_lines * line_h + pad_y * 2
                
                # Check drawing height if matched drawing
                d = cell_drawings.get((r_idx, col_idx))
                if d and pg_path:
                    try:
                        bbox = d.get("bbox")
                        cw_span = sum(col_widths[col_idx : col_idx + c_span])
                        
                        if r_idx == 0:
                            cropped = _crop_drawing(pg_path, bbox, padding=0.0, pad_px=15, pad_bottom=2)
                        else:
                            cropped = _crop_drawing(pg_path, bbox, padding=0.0, pad_px=20, pad_top=-12)
                        if cropped:
                            with Image.open(cropped) as ci:
                                nat_w, nat_h = ci.size
                            aspect = nat_h / nat_w if nat_w > 0 else 0.5
                            
                            # Scale signature based on bbox width if available
                            if bbox and len(bbox) == 4:
                                x1, y1, x2, y2 = [float(v) for v in bbox]
                                draw_w = usable_w * (x2 - x1)
                                draw_w = max(40.0, min(cw_span - pad_x * 2, draw_w))
                            else:
                                draw_w = cw_span - pad_x * 2
                            
                            draw_h = draw_w * aspect
                            if draw_h < 40:
                                draw_h = 40
                            
                            # If cell has text and drawing, sum their heights
                            cell_txt_val = table_data[r_idx][col_idx].strip() if col_idx < len(table_data[r_idx]) else ""
                            if cell_txt_val:
                                req_h = (max_lines * line_h) + draw_h + pad_y * 3
                            else:
                                req_h = max(req_h, draw_h + pad_y * 2)
                    except Exception:
                        pass
                
                # Distribute required height across all spanned rows
                current_span_h = sum(row_heights[r_idx : r_idx + r_span])
                if current_span_h < req_h:
                    diff = req_h - current_span_h
                    add_per_row = diff / r_span
                    for rx in range(r_idx, r_idx + r_span):
                        row_heights[rx] += add_per_row
        
    total_table_height = sum(row_heights)
    
    # Simple page wrap check
    if y - total_table_height < margin_pts:
        c.showPage()
        y = page_h - margin_pts

    # Draw grid and text
    c.setLineWidth(0.5)
    
    # Calculate top y for each row
    y_tops = []
    current_y = y
    for h in row_heights:
        y_tops.append(current_y)
        current_y -= h
    y_tops.append(current_y) # Bottom of last row
    
    for row_idx in range(rows):
        curr_x = x_start
        y_row_top = y_tops[row_idx]
        
        for col_idx in range(cols):
            cw_single = col_widths[col_idx]
            
            if not covered[row_idx][col_idx]:
                r_span, c_span = spans[row_idx][col_idx]
                cw_span = sum(col_widths[col_idx : col_idx + c_span])
                
                # Determine cell bottom y-coordinate based on row span
                span_bottom_idx = row_idx + r_span
                y_cell_bottom = y_tops[span_bottom_idx]
                h_cell = y_row_top - y_cell_bottom
                
                # Draw cell border
                if not borderless:
                    c.rect(curr_x, y_cell_bottom, cw_span, h_cell)
                
                # Check if there is a matched drawing for this cell
                d = cell_drawings.get((row_idx, col_idx))
                if d and pg_path:
                    try:
                        if row_idx == 0:
                            cropped = _crop_drawing(pg_path, d.get("bbox"), padding=0.0, pad_px=15, pad_bottom=2)
                        else:
                            cropped = _crop_drawing(pg_path, d.get("bbox"), padding=0.0, pad_px=20, pad_top=-12)
                        if cropped:
                            with Image.open(cropped) as ci:
                                nat_w, nat_h = ci.size
                            aspect = nat_h / nat_w if nat_w > 0 else 0.5
                            draw_w = cw_span - pad_x * 2
                            draw_h = draw_w * aspect
                            # Ensure minimum visual height
                            if draw_h < 50:
                                draw_h = 50
                                draw_w = draw_h / aspect
                            # Fit within cell height
                            if draw_h > h_cell - pad_y * 2:
                                draw_h = h_cell - pad_y * 2
                                draw_w = draw_h / aspect
                            # Center horizontally; place in upper portion of cell
                            img_x = curr_x + (cw_span - draw_w) / 2.0
                            img_y = y_cell_bottom + h_cell - pad_y - draw_h
                            c.drawImage(ImageReader(cropped), img_x, img_y, width=draw_w, height=draw_h,
                                        preserveAspectRatio=True, mask='auto')
                            
                            # Draw text in the lower portion of the cell if present
                            txt_val = table_data[row_idx][col_idx].strip() if col_idx < len(table_data[row_idx]) else ""
                            if txt_val:
                                c.setFont(font_name, fsize)
                                tw = c.stringWidth(txt_val, font_name, fsize)
                                text_x = curr_x + (cw_span - tw) / 2.0
                                text_y = y_cell_bottom + pad_y
                                c.drawString(text_x, text_y, txt_val)
                        else:
                            # Fallback if crop failed
                            txt_val = table_data[row_idx][col_idx].strip() if col_idx < len(table_data[row_idx]) else ""
                            if txt_val:
                                c.setFont(font_name, fsize)
                                c.drawString(curr_x + pad_x, y_cell_bottom + pad_y, txt_val)
                    except Exception as e:
                        print(f"  [PDF] Cell drawing failed: {e}")
                else:
                    # Draw wrapped lines inside the cell
                    f = font_bold if row_idx == 0 else font_name
                    c.setFont(f, fsize)
                    
                    lines = wrapped_data[row_idx][col_idx]
                    num_lines = len(lines)
                    total_text_h = num_lines * line_h
                    
                    # Vertically center text
                    start_offset_y = (h_cell - total_text_h) / 2.0
                    y_text = y_row_top - start_offset_y - fsize
                    
                    # Determine column alignment from col_alignments
                    col_aligns = (table_el.get("col_alignments") or []) if table_el else []
                    col_align = col_aligns[col_idx] if col_idx < len(col_aligns) else "left"
                    
                    for line in lines:
                        tw = c.stringWidth(line, f, fsize)
                        if row_idx == 0 or col_align == "center":
                            text_x = curr_x + (cw_span - tw) / 2.0
                        elif col_align == "right":
                            text_x = curr_x + cw_span - tw - pad_x
                        else:
                            text_x = curr_x + pad_x
                            
                        c.drawString(text_x, y_text, line)
                        y_text -= line_h
                    
            curr_x += cw_single
            
    return y_tops[-1] - 10


def _pdf_place_drawing_inline(c, image_path, bbox, y, page_w, margin_pts, page_h, line_spacing=1.15):
    """
    Place a drawing INLINE in the PDF text flow at the current y position.
    Size is derived from the actual cropped pixel dimensions (not raw bbox fractions),
    so large/small bboxes don't distort the drawing size.
    Returns the height consumed (in points) so the caller can advance y correctly.
    """
    if not bbox or len(bbox) != 4:
        return 0

    cropped = _crop_drawing(image_path, bbox)
    if not cropped:
        return 0

    try:
        content_w = page_w - 2.0 * margin_pts

        x1n, y1n, x2n, y2n = [float(v) for v in bbox]

        # Use actual pixel dimensions to derive aspect ratio — never trust raw bbox size
        with Image.open(cropped) as ci:
            nat_w, nat_h = ci.size
        aspect = nat_h / nat_w if nat_w > 0 else 1.0

        # Size derived from relative bbox width fraction of content area
        bbox_w_frac = x2n - x1n
        draw_w = content_w * bbox_w_frac
        draw_w = max(40.0, min(content_w, draw_w))
        draw_h = draw_w * aspect

        # Cap height at 60% of usable page height
        max_h = (page_h - 2.0 * margin_pts) * 0.60
        if draw_h > max_h:
            draw_h = max_h
            draw_w = draw_h / aspect if aspect > 0 else draw_w

        # Align horizontally based on original horizontal position in source image
        cx_norm = (x1n + x2n) / 2.0
        if cx_norm < 0.35:
            x = margin_pts
        elif cx_norm > 0.65:
            x = margin_pts + content_w - draw_w
        else:
            x = margin_pts + (content_w - draw_w) / 2.0

        # Check if there's enough room on the current page; if not, start new page
        if y - draw_h < margin_pts:
            c.showPage()
            y = page_h - margin_pts

        # ReportLab y is from bottom; we track y from top
        y_bottom = y - draw_h

        ir = ImageReader(cropped)
        c.drawImage(ir, x, y_bottom, width=draw_w, height=draw_h,
                    preserveAspectRatio=True, mask='auto')

        print(f"  [PDF] Inline drawing: x={x:.1f} y_btm={y_bottom:.1f} {draw_w:.1f}×{draw_h:.1f}pts (aspect={aspect:.2f})")
        return draw_h + 14   # height consumed + bottom padding
    except Exception as e:
        print(f"  [PDF] Inline drawing failed: {e}")
        return 0


# ═══════════════════════════════════════════════════
#  CONVERT ENDPOINT
# ═══════════════════════════════════════════════════
# ── API Status endpoint ──────────────────────────
@app.get("/api-status")
def api_status():
    """Return current provider health and which keys are configured."""
    return JSONResponse({
        "last_provider": _last_provider,
        "providers": {
            "gemini":     {
                "configured": bool(_GEMINI_KEYS),
                "key_count":  len(_GEMINI_KEYS),
                "status":     _provider_status.get("gemini", "unknown"),
            },
            "groq":       {
                "configured": bool(GROQ_API_KEY),
                "status":     _provider_status.get("groq", "not_configured"),
            },
            "openrouter": {
                "configured": bool(OPENROUTER_API_KEY),
                "status":     _provider_status.get("openrouter", "not_configured"),
            },
        }
    })


def pre_scan_match_drawings(elements, page_w, page_h, margin_pts):
    page_drawings = [item for item in elements if item.get("type") == "drawing"]
    y_est = page_h - margin_pts
    for el in elements:
        etype = el.get("type", "text")
        if etype == "blank_line":
            y_est -= 8
        elif etype == "arrow":
            y_est -= 30
        elif etype == "bracket":
            y_est -= 44
        elif etype == "drawing":
            # If not already consumed, it will consume space in the flow
            if not el.get("consumed", False):
                y_est -= 60
        elif etype == "table":
            table_data = el.get("data", [])
            if table_data:
                rows = len(table_data)
                cols = max(len(r) for r in table_data) if table_data else 0
                col_widths_est = [ (page_w - 2 * margin_pts) / cols ] * cols
                
                for col_idx in range(cols):
                    col_left = margin_pts + sum(col_widths_est[:col_idx])
                    col_right = col_left + col_widths_est[col_idx]
                    
                    for r_idx in range(0, rows):
                        cell_txt = table_data[r_idx][col_idx] if col_idx < len(table_data[r_idx]) else ""
                        txt_clean = cell_txt.strip().lower().replace("[", "").replace("]", "").replace(" ", "")
                        is_empty = txt_clean == "" or "signature" in txt_clean or "drawing" in txt_clean or "image" in txt_clean or "stamp" in txt_clean
                        
                        if is_empty:
                            for d in page_drawings:
                                if d.get("consumed", False):
                                    continue
                                bbox = d.get("bbox")
                                if bbox and len(bbox) == 4:
                                    x1, y1, x2, y2 = bbox
                                    cx_norm = (x1 + x2) / 2.0
                                    cy_norm = (y1 + y2) / 2.0
                                    draw_x = cx_norm * page_w
                                    draw_y = (1.0 - cy_norm) * page_h
                                    
                                    # Use dynamic row height estimation to prevent vertical matching overlap
                                    row_h_est = 18 if rows > 8 else 30
                                    y_min = y_est - rows * row_h_est - 40
                                    y_max = y_est + 40
                                    
                                    if (col_left <= draw_x <= col_right) and (y_min <= draw_y <= y_max):
                                        d["consumed"] = True
                                        d["matched_cell"] = (r_idx, col_idx)
                                        d["matched_table_id"] = id(el)
                                        break
                # Advance y_est using the dynamic row height estimation
                row_h_est = 18 if rows > 8 else 25
                y_est -= rows * row_h_est + 20
        else:
            y_est -= 16


def _clean_json_response(resp_text: str) -> str:
    text = resp_text.strip()
    if text.startswith("```"):
        first_newline = text.find("\n")
        if first_newline != -1:
            text = text[first_newline:].strip()
        if text.endswith("```"):
            text = text[:-3].strip()
    return text


@app.post("/convert")
async def convert(
    file:              UploadFile = File(...),
    mode:              str   = Form("preserve"),
    outtype:           str   = Form("docx"),
    auto_format:       str   = Form("true"),
    # Manual params
    font_family:       str   = Form("Calibri"),
    font_size:         float = Form(12.0),
    line_spacing:      float = Form(1.15),
    para_spacing:      float = Form(8.0),
    first_line_indent: float = Form(0.0),
    page_margin:       float = Form(2.54),
    text_align:        str   = Form("left"),
    text_bold:         str   = Form("false"),
    text_italic:       str   = Form("false"),
    text_underline:    str   = Form("false"),
):
    uid        = str(uuid.uuid4())
    input_path = f"{UPLOAD_DIR}/{uid}_{file.filename}"

    with open(input_path, "wb") as f:
        f.write(await file.read())

    use_auto = auto_format.lower() == "true"

    if input_path.lower().endswith(".pdf"):
        pages = pdf_to_images(input_path)
    else:
        pages = [input_path]

    # ═══════════════════════════
    #  WORD (.docx)
    # ═══════════════════════════
    if outtype == "docx":
        doc = Document()

        if use_auto:
            all_data = [run_ocr_auto(p) for p in pages]
            for elements, _ in all_data:
                _sanitize_elements(elements)
            has_table = any(any(el.get("type") == "table" for el in elements) for elements, _ in all_data)
            orig_margin = all_data[0][1].get("page_margin_cm", 2.54)
            margin_cm = 1.5 if has_table else orig_margin
            _set_margins(doc, margin_cm)

            for i, (elements, meta) in enumerate(all_data):
                if i > 0:
                    doc.add_page_break()
                ls = meta.get("line_spacing", 1.15)
                pm = margin_cm
                page_drawings = [item for item in elements if item.get("type") == "drawing"]
                table_idx = 0
                total_tables = sum(1 for item in elements if item.get("type") == "table")
                pre_scan_match_drawings(elements, page_w=595.27, page_h=841.89, margin_pts=pm*PT_PER_CM)

                for el in elements:
                    etype = el.get("type", "text")

                    if etype == "blank_line":
                        para = doc.add_paragraph()
                        para.paragraph_format.space_after = Pt(4)

                    elif etype == "drawing":
                        if el.get("consumed", False):
                            continue
                        # Inline image — sits naturally in the text flow
                        _add_inline_picture_docx(
                            doc,
                            img_path=pages[i],
                            bbox=el.get("bbox"),
                            page_margin_cm=pm,
                        )

                    elif etype == "arrow":
                        # Simple arrow — render as a large Unicode character, no image
                        arrow_char = el.get("arrow_char", "⟶")
                        para = doc.add_paragraph()
                        para.paragraph_format.alignment    = WD_ALIGN_PARAGRAPH.CENTER
                        para.paragraph_format.space_before = Pt(6)
                        para.paragraph_format.space_after  = Pt(6)
                        run = para.add_run(arrow_char)
                        run.font.name = "Segoe UI Symbol"
                        run.font.size = Pt(24)
                        run.font.bold = False
                        print(f"  [DOCX] Arrow element: '{arrow_char}'")

                    elif etype == "bracket":
                        # Grouping bracket — render as a tall bold Unicode char, no image
                        brk_char  = el.get("bracket_char", "[")
                        brk_side  = el.get("bracket_side", "left")
                        brk_style = el.get("bracket_style", "square")
                        # Choose alignment: left-bracket aligns left, right-bracket aligns right
                        align_val = (WD_ALIGN_PARAGRAPH.LEFT
                                     if brk_side == "left"
                                     else WD_ALIGN_PARAGRAPH.RIGHT)
                        para = doc.add_paragraph()
                        para.paragraph_format.alignment    = align_val
                        para.paragraph_format.space_before = Pt(4)
                        para.paragraph_format.space_after  = Pt(4)
                        run = para.add_run(brk_char)
                        # Use a larger font so the bracket is visually tall
                        run.font.name = "Segoe UI Symbol"
                        run.font.size = Pt(36)
                        run.font.bold = (brk_style == "plain")  # bold for plain brackets
                        print(f"  [DOCX] Bracket element: '{brk_char}' ({brk_style}, {brk_side})")

                    elif etype == "table":
                        # Native DOCX Table
                        table_data = el.get("data", [])
                        if table_data:
                            rows = len(table_data)
                            cols = max((len(r) for r in table_data), default=0)
                            if cols > 0:
                                table = doc.add_table(rows=rows, cols=cols)
                                if el.get("borderless", False):
                                    table.style = 'Normal Table'
                                else:
                                    table.style = 'Table Grid'
                                
                                # Set minimum row height for all rows to match original spacing
                                row_height_cm = 0.85  # approx 1 notebook line height
                                for row in table.rows:
                                    _set_row_height(row, row_height_cm)
                                
                                # Make a copy of table_data so we can safely update text
                                table_data = [list(r) for r in table_data]
                                
                                # Pre-process DATE column (index 0) to combine date and day of week
                                # Handles both old split-row format AND new \n-separator format from Gemini.
                                DAYS_OF_WEEK = {"monday", "tuesday", "wednesday", "thursday", "friday", "saturday", "sunday"}
                                for r_idx in range(1, rows):
                                    if len(table_data[r_idx]) == 0:
                                        continue
                                    val_curr = table_data[r_idx][0].strip()
                                    # Case 1: Gemini already combined with \n — normalize the separator
                                    if '\n' in val_curr:
                                        table_data[r_idx][0] = val_curr.replace('\n', ' ')
                                        continue
                                    # Case 2: Old format — day-of-week is on its own row
                                    if val_curr.lower() in DAYS_OF_WEEK:
                                        prev_r = r_idx - 1
                                        while prev_r >= 0 and (len(table_data[prev_r]) == 0 or table_data[prev_r][0].strip() == ""):
                                            prev_r -= 1
                                        if prev_r >= 0:
                                            val_prev = table_data[prev_r][0].strip()
                                            if val_curr.lower() not in val_prev.lower():
                                                table_data[prev_r][0] = (val_prev + " " + val_curr).strip()
                                                table_data[r_idx][0] = ""
                                
                                # 1a. Detect vertical cell spans (rowspans)
                                spans = [[(1, 1) for _ in range(cols)] for _ in range(rows)]
                                covered = [[False for _ in range(cols)] for _ in range(rows)]
                                ref_table = [list(r) for r in table_data]
                                
                                for col_idx in range(cols):
                                    # Find non-empty rows for this column (excluding header)
                                    non_empty_rows = []
                                    for r_idx in range(1, rows):
                                        cell_val = ref_table[r_idx][col_idx] if col_idx < len(ref_table[r_idx]) else ""
                                        if cell_val.strip() != "":
                                            non_empty_rows.append(r_idx)
                                            
                                    if not non_empty_rows:
                                        continue
                                        
                                    assignments = {}
                                    for r_idx in range(1, rows):
                                        cell_val = ref_table[r_idx][col_idx] if col_idx < len(ref_table[r_idx]) else ""
                                        if r_idx in non_empty_rows:
                                            assignments[r_idx] = r_idx
                                        else:
                                            # Find nearest above and below in non_empty_rows
                                            r_above = None
                                            for rx in reversed(non_empty_rows):
                                                if rx < r_idx:
                                                    r_above = rx
                                                    break
                                            r_below = None
                                            for rx in non_empty_rows:
                                                if rx > r_idx:
                                                    r_below = rx
                                                    break
                                                    
                                            if r_above is not None and r_below is not None:
                                                # Calculate similarity score based on compatibility with other columns
                                                score_above = 0
                                                score_below = 0
                                                for c in range(cols):
                                                    if c == col_idx:
                                                        continue
                                                    
                                                    val_r = ref_table[r_idx][c].strip().lower() if c < len(ref_table[r_idx]) else ""
                                                    val_above = ref_table[r_above][c].strip().lower() if c < len(ref_table[r_above]) else ""
                                                    val_below = ref_table[r_below][c].strip().lower() if c < len(ref_table[r_below]) else ""
                                                    
                                                    if val_r != "":
                                                        if val_r == val_above:
                                                            score_above += 1
                                                        if val_r == val_below:
                                                            score_below += 1
                                                
                                                if score_below > score_above:
                                                    assignments[r_idx] = r_below
                                                else:
                                                    assignments[r_idx] = r_above
                                            elif r_above is not None:
                                                assignments[r_idx] = r_above
                                            elif r_below is not None:
                                                assignments[r_idx] = r_below
                                                
                                    # Group by parent row to find spans
                                    groups = {}
                                    for r_idx, parent in assignments.items():
                                        groups.setdefault(parent, []).append(r_idx)
                                        
                                    for parent, row_list in groups.items():
                                        row_list.sort()
                                        start_r = row_list[0]
                                        span = len(row_list)
                                        
                                        spans[start_r][col_idx] = (span, 1)
                                        parent_text = ref_table[parent][col_idx] if col_idx < len(ref_table[parent]) else ""
                                        
                                        # Guard: row may be shorter than col_idx in a jagged OCR table
                                        if col_idx < len(table_data[start_r]):
                                            table_data[start_r][col_idx] = parent_text
                                        for rx in row_list:
                                            if rx != start_r:
                                                covered[rx][col_idx] = True
                                                if col_idx < len(table_data[rx]):
                                                    table_data[rx][col_idx] = ""

                                # Match drawings to cells using pre-scan assignments
                                cell_drawings = {}
                                if page_drawings:
                                    for d in page_drawings:
                                        if d.get("matched_table_id") == id(el):
                                            cell_drawings[d.get("matched_cell")] = d
                                            d["consumed"] = True

                                # 1b. Detect horizontal cell spans (colspans) on top of vertical spans
                                for r_idx in range(rows):
                                    col_idx = 0
                                    while col_idx < cols:
                                        if not covered[r_idx][col_idx]:
                                            r_span, c_span = spans[r_idx][col_idx]
                                            next_c = col_idx + 1
                                            while next_c < cols:
                                                # Merge next_c if it is empty, not covered, not matched to a drawing, and has same r_span
                                                # Guard: row may be shorter than next_c in a jagged OCR table
                                                next_c_val = table_data[r_idx][next_c] if next_c < len(table_data[r_idx]) else None
                                                if (next_c_val is not None and
                                                    not covered[r_idx][next_c] and 
                                                    (next_c_val.strip() == "") and 
                                                    (r_idx, next_c) not in cell_drawings and
                                                    spans[r_idx][next_c][0] == r_span):
                                                    
                                                    # Merge next_c into col_idx for all rows in the vertical span
                                                    for rx in range(r_idx, r_idx + r_span):
                                                        covered[rx][next_c] = True
                                                        if next_c < len(table_data[rx]):
                                                            table_data[rx][next_c] = ""
                                                    
                                                    c_span += 1
                                                    spans[r_idx][col_idx] = (r_span, c_span)
                                                    next_c += 1
                                                else:
                                                    break
                                            col_idx += c_span
                                        else:
                                            col_idx += 1
                                
                                # 2. Populate and merge cells in Word
                                for r_idx in range(rows):
                                    for col_idx in range(cols):
                                        if not covered[r_idx][col_idx]:
                                            r_span, c_span = spans[r_idx][col_idx]
                                            cell = table.cell(r_idx, col_idx)
                                            cell_text = table_data[r_idx][col_idx] if col_idx < len(table_data[r_idx]) else ""
                                            
                                            # Perform merge if row span or col span is greater than 1
                                            if r_span > 1 or c_span > 1:
                                                target_cell = table.cell(r_idx + r_span - 1, col_idx + c_span - 1)
                                                cell = cell.merge(target_cell)
                                                
                                            # Check if there is a matched drawing for this cell
                                            d = cell_drawings.get((r_idx, col_idx))
                                            if d:
                                                try:
                                                    cropped = _crop_drawing(pages[i], d.get("bbox"), padding=0.0, pad_px=20)
                                                    if cropped:
                                                        # Scale signature based on bbox width if available
                                                        bbox = d.get("bbox")
                                                        if bbox and len(bbox) == 4:
                                                            x1, y1, x2, y2 = [float(v) for v in bbox]
                                                            img_w_cm = usable_w_cm * (x2 - x1)
                                                            img_w_cm = max(1.5, min(col_w_cm * 0.85, img_w_cm))
                                                        else:
                                                            img_w_cm = max(2.0, min(col_w_cm * 0.85, 5.0 * c_span))
                                                        
                                                        with Image.open(cropped) as ci:
                                                            nat_w, nat_h = ci.size
                                                        aspect = nat_h / nat_w if nat_w > 0 else 0.5
                                                        img_h_cm = img_w_cm * aspect
                                                        
                                                        p = cell.paragraphs[0]
                                                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                                        run = p.add_run()
                                                        run.add_picture(cropped, width=Cm(img_w_cm), height=Cm(img_h_cm))
                                                        
                                                        # Render the signee's name text below the signature image
                                                        if cell_text:
                                                            p_text = cell.add_paragraph()
                                                            p_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                                            p_text.paragraph_format.space_before = Pt(4)
                                                            p_text.paragraph_format.space_after = Pt(2)
                                                            run_text = p_text.add_run(cell_text)
                                                            run_text.font.name = font_family if font_family else "Calibri"
                                                            run_text.font.size = Pt(9)
                                                            run_text.font.bold = False
                                                except Exception as e:
                                                    print(f"  [DOCX] Cell drawing failed: {e}")
                                            else:
                                                # Multi-line aware cell text rendering
                                                cell_lines = cell_text.split("\n") if cell_text else [""]
                                                # Clear the default empty paragraph
                                                first_para = cell.paragraphs[0]
                                                
                                                # Set cell padding to add visual breathing room
                                                _set_cell_margins(cell, top_pt=6, bottom_pt=6, left_pt=6, right_pt=6)
                                                
                                                # Determine alignment for this column
                                                col_aligns = el.get("col_alignments") or []
                                                col_align_str = col_aligns[col_idx] if col_idx < len(col_aligns) else ("center" if r_idx == 0 else "left")
                                                col_align_enum = ALIGN_MAP.get(col_align_str, WD_ALIGN_PARAGRAPH.LEFT)
                                                
                                                cell_font_size = Pt(9.5) if rows > 5 else Pt(11)
                                                is_header_row = (r_idx == 0)
                                                
                                                for li, line_text in enumerate(cell_lines):
                                                    if li == 0:
                                                        para = first_para
                                                    else:
                                                        para = cell.add_paragraph()
                                                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header_row else col_align_enum
                                                    para.paragraph_format.space_before = Pt(2)
                                                    para.paragraph_format.space_after = Pt(2)
                                                    run = para.add_run(line_text)
                                                    run.font.name = font_family if font_family else "Calibri"
                                                    run.font.size = cell_font_size
                                                    run.font.bold = is_header_row
                                
                                # Add space after table
                                doc.add_paragraph().paragraph_format.space_after = Pt(8)
                                print(f"  [DOCX] Table element: {rows}x{cols} (with vertical merges)")
                                table_idx += 1

                    else:  # "text"
                        txt = el.get("text", "")
                        fsize = float(el.get("font_size_pt", 12))
                        sp_before = float(el.get("space_before_pt", 0))
                        sp_after = float(el.get("space_after_pt", 4))
                        
                        is_note, is_metadata = _classify_text_element(txt)
                        if is_note:
                            fsize = 8.5
                            sp_before = 1
                            sp_after = 2
                        elif is_metadata:
                            fsize = 9.5
                            sp_before = 1
                            sp_after = 2

                        para = doc.add_paragraph()
                        run  = para.add_run(txt)
                        _apply_run(
                            run, font_family,
                            fsize,
                            el.get("bold", False),
                            el.get("italic", False),
                            el.get("underline", False),
                        )
                        _apply_para(
                            para,
                            el.get("alignment", "left"),
                            ls,
                            sp_before,
                            sp_after,
                            el.get("first_line_indent_cm", 0.0),
                            el.get("left_indent_cm", 0.0),
                        )

        else:
            bold      = text_bold.lower()      == "true"
            italic    = text_italic.lower()    == "true"
            underline = text_underline.lower() == "true"
            _set_margins(doc, page_margin)

            for i, page in enumerate(pages):
                lines = [_sanitize_xml_text(l) for l in run_ocr_manual(page)]
                if mode == "preserve" and i > 0:
                    doc.add_page_break()
                if mode == "flow":
                    text = " ".join(lines)
                    if text.strip():
                        para = doc.add_paragraph()
                        run  = para.add_run(text)
                        _apply_run(run, font_family, font_size, bold, italic, underline)
                        _apply_para(para, text_align, line_spacing,
                                    0, _snap_spacing(para_spacing), first_line_indent)
                else:
                    for line in lines:
                        para = doc.add_paragraph()
                        run  = para.add_run(line)
                        _apply_run(run, font_family, font_size, bold, italic, underline)
                        _apply_para(para, text_align, line_spacing,
                                    0, _snap_spacing(para_spacing), first_line_indent)

        out_path = f"{OUTPUT_DIR}/{uid}.docx"
        doc.save(out_path)
        return FileResponse(out_path, filename="output.docx")

    # ═══════════════════════════
    #  PDF
    # ═══════════════════════════
    else:
        out_path = f"{OUTPUT_DIR}/{uid}.pdf"
        c = canvas.Canvas(out_path, pagesize=A4)
        page_w, page_h = A4

        if use_auto:
            for pg_path in pages:
                elements, meta = run_ocr_auto(pg_path)
                _sanitize_elements(elements)
                has_table = any(el.get("type") == "table" for el in elements)
                margin_cm = 1.5 if has_table else meta.get("page_margin_cm", 2.54)
                margin_pts = max(28, margin_cm * PT_PER_CM)
                usable_w   = page_w - 2 * margin_pts
                ls         = meta.get("line_spacing", 1.15)
                page_drawings = [item for item in elements if item.get("type") == "drawing"]
                pre_scan_match_drawings(elements, page_w, page_h, margin_pts)

                # Single pass: render text AND drawings in order, top to bottom
                y = page_h - margin_pts

                for el in elements:
                    etype = el.get("type", "text")

                    if etype == "blank_line":
                        y -= 8
                        if y < margin_pts:
                            c.showPage()
                            y = page_h - margin_pts
                        continue

                    if etype == "arrow":
                        # Simple arrow — render as large centred Unicode character
                        arrow_char  = el.get("arrow_char", "⟶")
                        fsize_arrow = 24
                        y -= 6
                        if y < margin_pts:
                            c.showPage()
                            y = page_h - margin_pts
                        tw = c.stringWidth(arrow_char, _PDF_FONT_NORMAL, fsize_arrow)
                        x_arrow = margin_pts + (usable_w - tw) / 2
                        c.setFont(_PDF_FONT_NORMAL, fsize_arrow)
                        c.drawString(x_arrow, y, arrow_char)
                        y -= fsize_arrow + 6
                        print(f"  [PDF] Arrow element: '{arrow_char}'")
                        continue

                    if etype == "bracket":
                        # Grouping bracket — render as tall Unicode character in the flow
                        brk_char  = el.get("bracket_char", "[")
                        brk_side  = el.get("bracket_side", "left")
                        fsize_brk = 40   # tall bracket
                        y -= 4
                        if y < margin_pts:
                            c.showPage()
                            y = page_h - margin_pts
                        tw = c.stringWidth(brk_char, _PDF_FONT_BOLD, fsize_brk)
                        # Left bracket aligns left, right bracket aligns right
                        if brk_side == "right":
                            x_brk = margin_pts + usable_w - tw
                        else:
                            x_brk = margin_pts
                        c.setFont(_PDF_FONT_BOLD, fsize_brk)
                        c.drawString(x_brk, y, brk_char)
                        y -= fsize_brk + 4
                        print(f"  [PDF] Bracket element: '{brk_char}' ({brk_side})")
                        continue

                    if etype == "table":
                        table_data = el.get("data", [])
                        y -= 6
                        y = _pdf_draw_table(c, table_data, y, margin_pts, page_w, page_h, borderless=el.get("borderless", False), drawings=page_drawings, pg_path=pg_path, table_el=el)
                        print(f"  [PDF] Table element")
                        continue

                    if etype == "drawing":
                        if el.get("consumed", False):
                            continue
                        # Place drawing inline at current y, advance y by height consumed
                        consumed = _pdf_place_drawing_inline(
                            c, pg_path, el.get("bbox"), y,
                            page_w, margin_pts, page_h, ls
                        )
                        y -= consumed
                        continue

                    # text element — with word-wrap for long lines
                    txt       = el.get("text", "")
                    fsize     = float(el.get("font_size_pt", 12))
                    bold      = el.get("bold", False)
                    italic    = el.get("italic", False)
                    underline = el.get("underline", False)
                    align     = el.get("alignment", "left")
                    indent    = el.get("first_line_indent_cm", 0.0) * PT_PER_CM
                    left_i    = el.get("left_indent_cm", 0.0) * PT_PER_CM
                    sp_before = float(el.get("space_before_pt", 0))
                    sp_after  = float(el.get("space_after_pt", 4))
                    
                    # Dynamic compact layout for timetables and dense documents
                    is_note, is_metadata = _classify_text_element(txt)
                    if is_note:
                        fsize = 8.5
                        sp_before = 1
                        sp_after = 2
                        leading = fsize * 1.1
                    elif is_metadata:
                        fsize = 9.5
                        sp_before = 1
                        sp_after = 2
                        leading = fsize * 1.15
                    else:
                        leading   = fsize * ls

                    font_name = _rl_font(bold, italic)

                    y -= sp_before
                    if y < margin_pts:
                        c.showPage()
                        y = page_h - margin_pts

                    # ── Word-wrap long lines to fit page width ──
                    line_x_base = margin_pts + indent + left_i
                    avail_w = usable_w - indent - left_i

                    # Split into wrapped sub-lines, handling newlines (\n) first
                    wrapped_lines = []
                    raw_lines = str(txt).split("\n")
                    for raw_line in raw_lines:
                        words = raw_line.split(" ")
                        cur_line = ""
                        for word in words:
                            test = (cur_line + " " + word).strip()
                            if c.stringWidth(test, font_name, fsize) <= avail_w:
                                cur_line = test
                            else:
                                if cur_line:
                                    wrapped_lines.append(cur_line)
                                cur_line = word
                        if cur_line:
                            wrapped_lines.append(cur_line)
                    if not wrapped_lines:
                        wrapped_lines = [txt]

                    for li, wline in enumerate(wrapped_lines):
                        if y < margin_pts:
                            c.showPage()
                            y = page_h - margin_pts

                        tw = c.stringWidth(wline, font_name, fsize)
                        if align == "center":
                            x = margin_pts + (usable_w - tw) / 2
                        elif align == "right":
                            x = margin_pts + usable_w - tw
                        else:
                            x = line_x_base

                        _pdf_draw_text(c, wline, x, y, font_name, fsize, underline)
                        y -= leading

                    y -= sp_after

                    if y < margin_pts:
                        c.showPage()
                        y = page_h - margin_pts

                c.showPage()

        else:
            bold       = text_bold.lower()      == "true"
            italic     = text_italic.lower()    == "true"
            underline  = text_underline.lower() == "true"
            margin_pts = float(page_margin) * PT_PER_CM
            usable_w   = page_w - 2 * margin_pts
            leading    = float(font_size) * float(line_spacing)
            indent_pts = float(first_line_indent) * PT_PER_CM
            font_name  = _rl_font(bold, italic)
            sp_after   = _snap_spacing(para_spacing)

            for pg_path in pages:
                lines = [_sanitize_xml_text(l) for l in run_ocr_manual(pg_path)]
                y     = page_h - margin_pts

                if mode == "flow":
                    text  = " ".join(lines)
                    words = text.split(" ")
                    line_str = ""
                    for w in words:
                        test = line_str + w + " "
                        if c.stringWidth(test, font_name, font_size) > usable_w:
                            _pdf_draw_text(c, line_str.strip(), margin_pts, y, font_name, font_size, underline)
                            y -= leading
                            if y < margin_pts:
                                c.showPage(); y = page_h - margin_pts
                            line_str = w + " "
                        else:
                            line_str = test
                    if line_str.strip():
                        _pdf_draw_text(c, line_str.strip(), margin_pts, y, font_name, font_size, underline)
                        y -= leading + sp_after
                else:
                    for line in lines:
                        x = margin_pts + indent_pts
                        if text_align == "center":
                            tw = c.stringWidth(line, font_name, font_size)
                            x  = margin_pts + (usable_w - tw) / 2
                        elif text_align == "right":
                            tw = c.stringWidth(line, font_name, font_size)
                            x  = margin_pts + usable_w - tw
                        if y < margin_pts:
                            c.showPage(); y = page_h - margin_pts
                        _pdf_draw_text(c, line, x, y, font_name, font_size, underline)
                        y -= leading + sp_after * 0.75

                c.showPage()

        c.save()
        return FileResponse(out_path, filename="output.pdf")